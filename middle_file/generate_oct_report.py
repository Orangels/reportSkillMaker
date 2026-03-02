#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成2025年10月警情分析报告
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 读取提取的数据
with open('/home/orangels/xm_dev/ls_dev/reportSkillMaker/middle_file/extracted_data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 创建文档
doc = Document()

# 设置文档默认字体
doc.styles['Normal'].font.name = '仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.styles['Normal'].font.size = Pt(16)
doc.styles['Normal'].paragraph_format.line_spacing = Pt(28)
doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0.5)

def add_title_paragraph(text, font_name='仿宋', font_size=16, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """添加标题段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_content_paragraph(text, bold_parts=None):
    """添加正文段落，支持部分加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.first_line_indent = Inches(0.5)

    if bold_parts:
        # 处理需要加粗的部分
        last_end = 0
        for bold_text in bold_parts:
            start = text.find(bold_text, last_end)
            if start != -1:
                # 添加加粗前的文本
                if start > last_end:
                    run = p.add_run(text[last_end:start])
                    run.font.name = '仿宋'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

                # 添加加粗文本
                run = p.add_run(bold_text)
                run.font.name = '仿宋'
                run.font.size = Pt(16)
                run.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

                last_end = start + len(bold_text)

        # 添加剩余文本
        if last_end < len(text):
            run = p.add_run(text[last_end:])
            run.font.name = '仿宋'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    else:
        run = p.add_run(text)
        run.font.name = '仿宋'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    return p

def add_red_line():
    """添加红色分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run('━' * 30)
    run.font.color.rgb = RGBColor(255, 0, 0)
    return p

# 1. 发文单位
add_title_paragraph('临高县公安局情报指挥中心', font_name='仿宋', font_size=16)

# 2. 红色分隔线
add_red_line()

# 3. 报告标题
add_title_paragraph('关于10月份警情分析研判报告', font_name='方正小标宋简体', font_size=22)

# 4. 空行
doc.add_paragraph()

# 5. 一、整体情况
add_title_paragraph('一、整体情况', font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

# 整体情况内容
overall = data['整体情况']
text = f"10月1日至31日，我局共接报有效警情{overall['有效警情总量']['本期']}起，环比上升{overall['有效警情总量']['环比']}%。"
bold_parts = ['有效警情']
add_content_paragraph(text, bold_parts)

text = f"其中刑事警情{overall['刑事警情']['本期']}起，环比上升{overall['刑事警情']['环比']}%；"
text += f"治安警情{overall['治安警情']['本期']}起，环比上升{overall['治安警情']['环比']}%；"
text += f"交通警情{overall['交通警情']['本期']}起，环比上升{overall['交通警情']['环比']}%；"
text += f"纠纷警情{overall['纠纷警情']['本期']}起，环比上升{overall['纠纷警情']['环比']}%；"
text += f"群众紧急求助{overall['群众紧急求助']['本期']}起，环比上升{overall['群众紧急求助']['环比']}%；"
text += f"其他警情{overall['其他警情']['本期']}起，环比上升{overall['其他警情']['环比']}%。"
bold_parts = ['刑事警情', '治安警情', '交通警情', '纠纷警情', '群众紧急求助', '其他警情']
add_content_paragraph(text, bold_parts)

# 6. 二、上升警情类别分布
add_title_paragraph('二、上升警情类别分布', font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

# （一）治安殴打他人警情分析
add_content_paragraph('（一）治安殴打他人警情分析', bold_parts=['（一）治安殴打他人警情分析'])

assault = data['治安殴打他人警情']
text = f"10月份，我局共接报治安殴打他人警情{assault['总量']['本期']}起，环比下降{abs(assault['总量']['环比'])}%。"
add_content_paragraph(text)

# 1.从发生原因分析
add_content_paragraph('1.从发生原因分析', bold_parts=['1.从发生原因分析'])
reasons = assault['发生原因分布']
if reasons:
    reason_list = sorted(reasons.items(), key=lambda x: x[1]['数量'], reverse=True)
    text = "主要原因包括"
    for i, (reason, info) in enumerate(reason_list):
        if i == 0:
            text += f"{reason}{info['数量']}起（占比{info['占比']}%）"
        elif i == len(reason_list) - 1:
            text += f"，{reason}{info['数量']}起（占比{info['占比']}%）"
        else:
            text += f"，其次{reason}{info['数量']}起（占比{info['占比']}%）"
    text += "。"
    add_content_paragraph(text)

# 2.从行为手段分析
add_content_paragraph('2.从行为手段分析', bold_parts=['2.从行为手段分析'])
methods = assault['行为手段分布']
method_list = sorted(methods.items(), key=lambda x: x[1]['数量'], reverse=True)
text = "从行为手段看，"
for i, (method, info) in enumerate(method_list):
    if i == 0:
        text += f"{method}{info['数量']}起（占比{info['占比']}%）"
    elif i == len(method_list) - 1:
        text += f"，{method}{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，其次{method}{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# 3.从涉刀警情分析
add_content_paragraph('3.从涉刀警情分析', bold_parts=['3.从涉刀警情分析'])
text = f"10月份涉刀警情{assault['涉刀警情']['本期']}起，环比上升{assault['涉刀警情']['环比']}%，需予以重点关注。"
add_content_paragraph(text)

# 4.从辖区分布分析
add_content_paragraph('4.从辖区分布分析', bold_parts=['4.从辖区分布分析'])

# （1）殴打警情
add_content_paragraph('（1）殴打警情', bold_parts=['（1）殴打警情'])
areas = assault['辖区分布']
area_list = sorted(areas.items(), key=lambda x: x[1]['数量'], reverse=True)[:3]
text = "主要发生在"
for i, (area, info) in enumerate(area_list):
    if i == 0:
        text += f"{area}派出所{info['数量']}起（占比{info['占比']}%）"
    elif i == len(area_list) - 1:
        text += f"，{area}派出所{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，其次{area}派出所{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# （2）涉刀警情
add_content_paragraph('（2）涉刀警情', bold_parts=['（2）涉刀警情'])
knife_areas = assault['涉刀警情辖区分布']
knife_list = sorted(knife_areas.items(), key=lambda x: x[1]['数量'], reverse=True)
text = "涉刀警情主要集中在"
for i, (area, info) in enumerate(knife_list):
    if i == 0:
        text += f"{area}派出所{info['数量']}起（占比{info['占比']}%）"
    elif i == len(knife_list) - 1:
        text += f"，{area}派出所{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，{area}派出所{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# 小结
text = f"小结：治安殴打他人警情虽环比下降{abs(assault['总量']['环比'])}%，但涉刀警情环比大幅上升{assault['涉刀警情']['环比']}%，此类警情特征突出。"
text += f"一是拳脚殴打仍是主要手段，占比{assault['行为手段分布']['拳脚殴打']['占比']}%，但刀具使用呈上升态势，需加强管控；"
text += f"二是临城西门派出所辖区警情高发，占比{assault['辖区分布']['临城西门']['占比']}%，且涉刀警情集中，需针对性强化巡防；"
text += "三是口角纠纷、邻里纠纷等矛盾易引发殴打行为，建议加强矛盾纠纷排查化解，从源头预防此类警情发生。"
bold_parts = ['小结：', '一是', '二是', '三是']
add_content_paragraph(text, bold_parts)

# （二）涉未成人警情分析
add_content_paragraph('（二）涉未成人警情分析', bold_parts=['（二）涉未成人警情分析'])

minor = data['涉未成人警情']
text = f"10月份，我局共接报涉未成人警情{minor['总量']['本期']}起，环比下降{abs(minor['总量']['环比'])}%。"
add_content_paragraph(text)

# 1.从警情类型分析
add_content_paragraph('1.从警情类型分析', bold_parts=['1.从警情类型分析'])
types = minor['警情类型分布']
type_list = sorted(types.items(), key=lambda x: x[1]['数量'], reverse=True)
text = "从警情类型看，"
for i, (ptype, info) in enumerate(type_list):
    if i == 0:
        text += f"{ptype}{info['数量']}起（占比{info['占比']}%）"
    elif i == len(type_list) - 1:
        text += f"，{ptype}{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，其次{ptype}{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# 2.从高发类型分析
add_content_paragraph('2.从高发类型分析', bold_parts=['2.从高发类型分析'])
high_freq = minor['高发类型']
text = "高发类型主要为"
items = sorted(high_freq.items(), key=lambda x: x[1], reverse=True)
for i, (htype, count) in enumerate(items):
    if i == 0:
        text += f"{htype}{count}起"
    elif i == len(items) - 1:
        text += f"、{htype}{count}起"
    else:
        text += f"、{htype}{count}起"
text += "。"
add_content_paragraph(text)

# 3.从敏感警情分析
add_content_paragraph('3.从敏感警情分析', bold_parts=['3.从敏感警情分析'])
text = f"10月份未接报涉未成人敏感警情。"
add_content_paragraph(text)

# 4.从辖区分布分析
add_content_paragraph('4.从辖区分布分析', bold_parts=['4.从辖区分布分析'])
minor_areas = minor['辖区分布']
minor_list = sorted(minor_areas.items(), key=lambda x: x[1]['数量'], reverse=True)[:3]
text = "主要发生在"
for i, (area, info) in enumerate(minor_list):
    if i == 0:
        text += f"{area}派出所{info['数量']}起（占比{info['占比']}%）"
    elif i == len(minor_list) - 1:
        text += f"，{area}派出所{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，其次{area}派出所{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# 小结
text = f"小结：涉未成人警情环比下降{abs(minor['总量']['环比'])}%，呈现良好态势。"
text += f"一是道路交通事故和殴打他人警情是主要类型，需加强未成年人交通安全教育和校园周边治安管控；"
text += f"二是临城西门、临城东门派出所辖区警情相对集中，占比分别为{minor['辖区分布']['临城西门']['占比']}%和{minor['辖区分布']['临城东门']['占比']}%，建议加强学校周边巡防；"
text += "三是本月未接报敏感警情，反映出未成年人保护工作成效显著，需持续巩固。"
bold_parts = ['小结：', '一是', '二是', '三是']
add_content_paragraph(text, bold_parts)

# （三）金牌港重点园区警情分析
add_content_paragraph('（三）金牌港重点园区警情分析', bold_parts=['（三）金牌港重点园区警情分析'])

park = data['金牌港重点园区警情']
text = f"10月份，金牌港重点园区共接报警情{park['总量']['本期']}起，环比下降{abs(park['总量']['环比'])}%。"
add_content_paragraph(text)

# 各类警情分布
park_types = park['各类警情分布']
text = "其中"
items = []
for ptype, info in park_types.items():
    if info['环比'] > 0:
        items.append(f"{ptype}{info['本期']}起，环比上升{info['环比']}%")
    elif info['环比'] < 0:
        items.append(f"{ptype}{info['本期']}起，环比下降{abs(info['环比'])}%")
    else:
        items.append(f"{ptype}{info['本期']}起")

text += "；".join(items) + "。"
add_content_paragraph(text)

# 纠纷警情细分
if park['纠纷警情细分']:
    text = "纠纷警情主要涉及"
    items = [f"{reason}{count}起" for reason, count in park['纠纷警情细分'].items()]
    text += "、".join(items) + "。"
    add_content_paragraph(text)

# 小结
text = f"小结：金牌港重点园区警情环比下降{abs(park['总量']['环比'])}%，总体态势平稳。"
text += f"一是行政（治安）类警情实现零接报，反映出园区治安管控成效显著；"
text += f"二是道路交通类警情环比下降{abs(park_types['道路交通类警情']['环比'])}%，但仍占较大比重，需持续加强交通安全管理；"
text += f"三是群众紧急求助环比上升{park_types['群众紧急求助']['环比']}%，建议马袅海岸派出所加强园区服务保障，及时回应群众诉求。"
bold_parts = ['小结：', '一是', '二是', '三是']
add_content_paragraph(text, bold_parts)

# 7. 三、下降警情类型分布
add_title_paragraph('三、下降警情类型分布', font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

# （一）交通事故警情分析
add_content_paragraph('（一）交通事故警情分析', bold_parts=['（一）交通事故警情分析'])

traffic = data['交通事故警情']
text = f"10月份，我局共接报交通警情{traffic['交通警情总量']['本期']}起，环比上升{traffic['交通警情总量']['环比']}%；"
text += f"其中交通事故警情{traffic['交通事故警情']['本期']}起，环比上升{traffic['交通事故警情']['环比']}%。"
add_content_paragraph(text)

# 1.从警情类别分析
add_content_paragraph('1.从警情类别分析', bold_parts=['1.从警情类别分析'])
categories = traffic['警情类别分布']
cat_list = sorted(categories.items(), key=lambda x: x, reverse=True)
text = "从警情类别看，"
for i, (cat, count) in enumerate(cat_list):
    if i == 0:
        text += f"{cat}{count}起"
    elif i == len(cat_list) - 1:
        text += f"，{cat}{count}起"
    else:
        text += f"，{cat}{count}起"
text += "。"
add_content_paragraph(text)

# 2.从发案时间分析
add_content_paragraph('2.从发案时间分析', bold_parts=['2.从发案时间分析'])
time_dist = traffic['时段分布']
time_list = sorted(time_dist.items(), key=lambda x: x[1]['数量'], reverse=True)
text = "从时段分布看，"
for i, (time_range, info) in enumerate(time_list):
    if i == 0:
        text += f"{time_range}时段{info['数量']}起（占比{info['占比']}%）"
    elif i == len(time_list) - 1:
        text += f"，{time_range}时段{info['数量']}起（占比{info['占比']}%）"
    else:
        text += f"，{time_range}时段{info['数量']}起（占比{info['占比']}%）"
text += "。"
add_content_paragraph(text)

# 小结
text = f"小结：交通事故警情环比上升{traffic['交通事故警情']['环比']}%，呈现明显上升态势。"
text += f"一是道路交通事故占绝对主体，占比超过98%，需持续加强道路交通安全管理；"
text += f"二是12-18时时段为高发时段，占比{time_dist['12-18时']['占比']}%，建议加强该时段路面巡查和交通疏导；"
text += "三是交通事故警情大幅上升反映出道路交通安全形势严峻，建议交警部门加强重点路段、重点时段管控，严查交通违法行为，切实预防和减少交通事故发生。"
bold_parts = ['小结：', '一是', '二是', '三是']
add_content_paragraph(text, bold_parts)

# 8. 四、工作建议
add_title_paragraph('四、工作建议', font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

# （一）严管严控持续压降涉刀警情
add_content_paragraph('（一）严管严控持续压降涉刀警情', bold_parts=['（一）严管严控持续压降涉刀警情'])
text = "临城西门、新盈海岸、加来等派出所要针对涉刀警情上升态势，深入开展刀具管控专项行动，加强对五金店、农资店等重点场所的检查，严格落实刀具销售实名登记制度；"
text += "要加强街面巡逻，及时发现和制止携带管制刀具行为；"
text += "要强化矛盾纠纷排查化解，对可能引发涉刀警情的矛盾纠纷及时介入调处，从源头上预防涉刀警情发生，切实维护社会治安稳定。"
add_content_paragraph(text)

# （二）多措并举筑牢未成年人保护屏障
add_content_paragraph('（二）多措并举筑牢未成年人保护屏障', bold_parts=['（二）多措并举筑牢未成年人保护屏障'])
text = "各派出所要持续加强校园周边治安管控，在上下学时段加强巡逻，及时发现和处置涉未成年人警情；"
text += "要深入开展交通安全宣传教育，提高未成年人交通安全意识；"
text += "要加强与教育部门、学校的协作配合，建立健全涉未成年人警情快速处置机制，确保涉未成年人警情得到及时妥善处理。"
add_content_paragraph(text)

# （三）精准施策护航园区平安稳定
add_content_paragraph('（三）精准施策护航园区平安稳定', bold_parts=['（三）精准施策护航园区平安稳定'])
text = "马袅海岸派出所要紧盯金牌港重点园区警情变化，加强园区巡逻防控，提升见警率、管事率；"
text += "要主动对接园区企业，及时了解掌握企业生产经营情况和治安需求，提供精准服务保障；"
text += "要加强劳资纠纷、经济纠纷等矛盾排查化解，防止矛盾激化引发群体性事件，全力维护园区治安秩序稳定。"
add_content_paragraph(text)

# （四）强化路段时段管控遏制交通事故
add_content_paragraph('（四）强化路段时段管控遏制交通事故', bold_parts=['（四）强化路段时段管控遏制交通事故'])
text = "交警部门要针对交通事故警情上升态势，加强12-18时、18-24时等高发时段的路面管控，严查超速、酒驾、疲劳驾驶等交通违法行为；"
text += "要加强事故多发路段的隐患排查整治，完善交通安全设施；"
text += "要深入开展交通安全宣传教育，提高驾驶人和群众的交通安全意识，从源头上预防和减少交通事故发生。"
add_content_paragraph(text)

# 9. 落款
doc.add_paragraph()
add_title_paragraph('临高县公安局情报指挥中心', alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_title_paragraph('2025年11月5日', alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# 10. 抄送信息
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(28)
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('抄送：县委办、县政府办、县委政法委。')
run.font.name = '仿宋'
run.font.size = Pt(16)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(28)
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('抄报：局领导。')
run.font.name = '仿宋'
run.font.size = Pt(16)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(28)
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('临高县公安局情报指挥中心                2025年11月5日印发')
run.font.name = '仿宋'
run.font.size = Pt(16)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# 保存文档
output_path = '/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_2025年10月统计报告.docx'
doc.save(output_path)
print(f"报告已生成：{output_path}")
