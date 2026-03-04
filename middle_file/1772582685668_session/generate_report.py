#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成2025年10月警情分析研判报告
严格遵循模板分析格式规范，智能仿写内容
"""

import json
import os
from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

# ============================================================
# 路径配置
# ============================================================
SESSION_DIR = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/middle_file/1772582685668_session"
OUTPUT_DIR = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/output"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "output_2025年10月统计报告.docx")
DATA_FILE = os.path.join(SESSION_DIR, "extracted_data.json")

# 加载数据
with open(DATA_FILE, 'r', encoding='utf-8') as f:
    data = json.load(f)

# ============================================================
# 格式工具函数
# ============================================================

def set_line_spacing_exact(paragraph, twips_value):
    """设置固定行距 - 直接操作XML，使用twips原始值"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(
            f'<w:spacing {nsdecls("w")} w:line="{twips_value}" w:lineRule="exact"/>'
        )
        pPr.append(spacing)
    else:
        spacing.set(qn('w:line'), str(twips_value))
        spacing.set(qn('w:lineRule'), 'exact')


def set_first_line_indent(paragraph, twips_value=640, chars=200):
    """设置首行缩进 - 使用twips原始值"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLine="{twips_value}" w:firstLineChars="{chars}"/>'
        )
        pPr.append(ind)
    else:
        ind.set(qn('w:firstLine'), str(twips_value))
        ind.set(qn('w:firstLineChars'), str(chars))


def set_left_indent(paragraph, twips_value):
    """设置左侧缩进"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:left="{twips_value}"/>'
        )
        pPr.append(ind)
    else:
        ind.set(qn('w:left'), str(twips_value))


def set_paragraph_properties(paragraph):
    """设置通用段落属性: widowControl=0, kinsoku, overflowPunct, snapToGrid"""
    pPr = paragraph._element.get_or_add_pPr()
    # widowControl = 0
    wc = pPr.find(qn('w:widowControl'))
    if wc is None:
        wc = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>')
        pPr.insert(0, wc)
    else:
        wc.set(qn('w:val'), '0')


def set_font_for_run(run, font_name, font_size_half_points, bold=False, color=None):
    """设置run的字体属性
    font_size_half_points: 半磅值，如32表示16磅
    """
    rPr = run._element.get_or_add_rPr()

    # 字体 - 同时设置ascii和eastAsia
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    # 字号 - 使用半磅值
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{font_size_half_points}"/>')
        rPr.append(sz)
    else:
        sz.set(qn('w:val'), str(font_size_half_points))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{font_size_half_points}"/>')
        rPr.append(szCs)
    else:
        szCs.set(qn('w:val'), str(font_size_half_points))

    # 加粗
    if bold:
        b = rPr.find(qn('w:b'))
        if b is None:
            b = parse_xml(f'<w:b {nsdecls("w")}/>')
            rPr.append(b)
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
            rPr.append(bCs)
    else:
        # 明确移除加粗
        b = rPr.find(qn('w:b'))
        if b is not None:
            rPr.remove(b)
        bCs = rPr.find(qn('w:bCs'))
        if bCs is not None:
            rPr.remove(bCs)

    # 颜色
    if color:
        c = rPr.find(qn('w:color'))
        if c is None:
            c = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
            rPr.append(c)
        else:
            c.set(qn('w:val'), color)


def add_run_with_format(paragraph, text, font_name, size_half_pt, bold=False, color=None):
    """添加格式化的run"""
    run = paragraph.add_run(text)
    set_font_for_run(run, font_name, size_half_pt, bold, color)
    return run


def create_body_paragraph(doc, align='both'):
    """创建标准正文段落，含基础格式"""
    p = doc.add_paragraph()
    set_line_spacing_exact(p, 560)
    set_first_line_indent(p, 640, 200)
    set_paragraph_properties(p)
    if align == 'both':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def create_level1_title(doc, text):
    """一级标题：黑体 16磅 两端对齐 首行缩进"""
    p = create_body_paragraph(doc, 'both')
    add_run_with_format(p, text, '黑体', 32, bold=False, color='000000')
    return p


def create_level2_title(doc, text):
    """二级标题：楷体 16磅 左对齐 首行缩进"""
    p = create_body_paragraph(doc, 'both')
    add_run_with_format(p, text, '楷体', 32, bold=False, color='000000')
    return p


def create_level3_paragraph(doc, title_text, body_text):
    """三级段落：标题仿宋16磅加粗，正文仿宋16磅不加粗"""
    p = create_body_paragraph(doc, 'both')
    set_first_line_indent(p, 643, 200)
    add_run_with_format(p, title_text, '仿宋', 32, bold=True, color='000000')
    add_run_with_format(p, body_text, '仿宋', 32, bold=False, color='000000')
    return p


def create_level4_paragraph(doc, title_text, body_text):
    """四级段落：标题仿宋16磅加粗，正文仿宋16磅不加粗"""
    p = create_body_paragraph(doc, 'both')
    set_first_line_indent(p, 643, 200)
    add_run_with_format(p, title_text, '仿宋', 32, bold=True, color='000000')
    add_run_with_format(p, body_text, '仿宋', 32, bold=False, color='000000')
    return p


def create_summary_paragraph(doc, body_text):
    """小结段落：'小结：'加粗，正文不加粗"""
    p = create_body_paragraph(doc, 'both')
    set_first_line_indent(p, 643, 200)
    add_run_with_format(p, '小结：', '仿宋', 32, bold=True, color='000000')
    add_run_with_format(p, body_text, '仿宋', 32, bold=False, color='000000')
    return p


def create_normal_paragraph(doc, text):
    """普通正文段落：仿宋 16磅"""
    p = create_body_paragraph(doc, 'both')
    add_run_with_format(p, text, '仿宋', 32, bold=False, color='000000')
    return p


def create_bold_keyword_paragraph(doc, segments):
    """
    创建包含加粗关键词的混合段落
    segments: [(text, bold), (text, bold), ...]
    """
    p = create_body_paragraph(doc, 'both')
    for text, bold in segments:
        add_run_with_format(p, text, '仿宋', 32, bold=bold, color='000000')
    return p


def add_red_line(doc):
    """添加红色分隔线（红头下方）"""
    p = doc.add_paragraph()
    set_line_spacing_exact(p, 560)
    pPr = p._element.get_or_add_pPr()
    # 下边框 - 红色粗线
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="thinThickSmallGap" w:sz="45" w:space="1" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_black_line_top(doc):
    """添加黑色上边框分隔线"""
    p = doc.add_paragraph()
    set_line_spacing_exact(p, 560)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="10" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_black_line_bottom(doc):
    """添加黑色下边框分隔线"""
    p = doc.add_paragraph()
    set_line_spacing_exact(p, 560)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def strip_prefix(name):
    """去掉派出所名称前缀 '临高'"""
    if name.startswith('临高'):
        return name[2:]
    return name


# ============================================================
# 创建文档
# ============================================================

doc = Document()

# ------ 页面设置 ------
section = doc.sections[0]
section.page_width = Emu(11906 * 635)   # 11906 twips -> EMU
section.page_height = Emu(16838 * 635)
section.top_margin = Emu(1962 * 635)
section.bottom_margin = Emu(1848 * 635)
section.left_margin = Emu(1587 * 635)
section.right_margin = Emu(1474 * 635)
section.header_distance = Emu(851 * 635)
section.footer_distance = Emu(992 * 635)

# 文档网格
sectPr = section._sectPr
docGrid = sectPr.find(qn('w:docGrid'))
if docGrid is None:
    docGrid = parse_xml(f'<w:docGrid {nsdecls("w")} w:type="lines" w:linePitch="312"/>')
    sectPr.append(docGrid)
else:
    docGrid.set(qn('w:type'), 'lines')
    docGrid.set(qn('w:linePitch'), '312')

# ============================================================
# 一、红头区域
# ============================================================

# 发文单位名称 - 方正小标宋简体 55磅 红色 居中
p_header = doc.add_paragraph()
p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing_exact(p_header, 560)
run_header = p_header.add_run('临高县公安局情报指挥中心')
set_font_for_run(run_header, '方正小标宋简体', 110, bold=False, color='FF0000')

# 红色分隔线
add_red_line(doc)

# ============================================================
# 二、主标题
# ============================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing_exact(p_title, 560)
run_title = p_title.add_run('关于10月份警情分析研判的报告')
set_font_for_run(run_title, '方正小标宋简体', 44, bold=False, color='000000')
# bCs only (模板中标题有bCs但无b)
rPr = run_title._element.get_or_add_rPr()
bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
rPr.append(bCs)

# ============================================================
# 三、正文内容
# ============================================================

# =========================
# 第一章：整体情况
# =========================
create_level1_title(doc, '一、整体情况')

# 整体情况段 - 包含加粗关键词
ov = data['整体情况']
cats = ov['各大类警情']

# 构建整体情况段落的segments
segments = []
segments.append(('10月1日至31日我局共接报', False))
segments.append(('有效警情', True))
segments.append((f'{ov["有效警情总数_本期"]}起（', False))
segments.append(('不含骚扰警情', True))
segments.append((f'{ov["骚扰警情数_本期"]}起），环比上升{ov["有效警情环比变化率"]}%。其中', False))

for i, cat in enumerate(cats):
    name = cat['类别']
    num = cat['本期数量']
    rate = cat['环比变化率']
    direction = '上升' if cat['环比方向'] == '上升' else '下降'
    segments.append((name, True))
    if i < len(cats) - 1:
        segments.append((f'{num}起，环比{direction}{rate}%；', False))
    else:
        segments.append((f'{num}起，环比{direction}{rate}%。', False))

p = create_body_paragraph(doc, 'both')
for text, bold in segments:
    add_run_with_format(p, text, '仿宋', 32, bold=bold, color='000000')


# =========================
# 第二章：上升警情类别分布
# =========================
create_level1_title(doc, '二、上升警情类别分布')

# ---- (一) 治安警情分析 ----
za = data['各类详细分析']['治安警情']
create_level2_title(doc, '（一）治安警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '治安警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{za["本期总量"]}起，环比上升{za["环比变化率"]}%。', '仿宋', 32, False, '000000')

# 1. 从警情类型分析
type_dist = za['反馈报警类型分布']
# 排序
type_sorted = sorted(type_dist, key=lambda x: x['本期数量'], reverse=True)
body_parts = '其中'
for i, t in enumerate(type_sorted[:4]):
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '环比持平')
    body_parts += f'{t["类型"]}{t["本期数量"]}起，{rate_text}'
    if i < 3:
        body_parts += '；'
    else:
        body_parts += '。'
# 补充其余
remaining = type_sorted[4:]
if remaining:
    body_parts += '另有'
    for i, t in enumerate(remaining):
        body_parts += f'{t["类型"]}{t["本期数量"]}起'
        if i < len(remaining) - 1:
            body_parts += '，'
        else:
            body_parts += '。'
create_level3_paragraph(doc, '1.从警情类型分析。', body_parts)

# 2. 从高发警情类型分析
fine_dist = za['反馈报警细类分布']
top_fine = fine_dist[:5]
body_text = f'治安警情中高发类型主要集中在{top_fine[0]["名称"]}{top_fine[0]["数量"]}起（占比{top_fine[0]["占比"]}%），其次{top_fine[1]["名称"]}{top_fine[1]["数量"]}起（占比{top_fine[1]["占比"]}%），{top_fine[2]["名称"]}{top_fine[2]["数量"]}起（占比{top_fine[2]["占比"]}%），{top_fine[3]["名称"]}{top_fine[3]["数量"]}起（占比{top_fine[3]["占比"]}%），{top_fine[4]["名称"]}{top_fine[4]["数量"]}起（占比{top_fine[4]["占比"]}%）。'
create_level3_paragraph(doc, '2.从高发警情类型分析。', body_text)

# 3. 从发案时段分析
ts = za['时段分布']
ts_sorted = sorted(ts, key=lambda x: x['数量'], reverse=True)
body_text = f'治安警情发案时段主要集中在{ts_sorted[0]["时段"]}{ts_sorted[0]["数量"]}起（占比{ts_sorted[0]["占比"]}%），其次{ts_sorted[1]["时段"]}{ts_sorted[1]["数量"]}起（占比{ts_sorted[1]["占比"]}%），{ts_sorted[2]["时段"]}{ts_sorted[2]["数量"]}起（占比{ts_sorted[2]["占比"]}%）。凌晨和夜间时段共发案{ts[0]["数量"]+ts[6]["数量"]}起，占治安警情总量的{round((ts[0]["数量"]+ts[6]["数量"])/za["本期总量"]*100, 1)}%，夜间治安问题较为突出。'
create_level3_paragraph(doc, '3.从发案时段分析。', body_text)

# 4. 从辖区分布分析
jq_dist = za['辖区分布']
top_jq = jq_dist[:5]
body_text = f'治安警情主要集中在{strip_prefix(top_jq[0]["名称"])}{top_jq[0]["数量"]}起（占比{top_jq[0]["占比"]}%），其次{strip_prefix(top_jq[1]["名称"])}{top_jq[1]["数量"]}起（占比{top_jq[1]["占比"]}%），{strip_prefix(top_jq[2]["名称"])}{top_jq[2]["数量"]}起（占比{top_jq[2]["占比"]}%），{strip_prefix(top_jq[3]["名称"])}{top_jq[3]["数量"]}起（占比{top_jq[3]["占比"]}%），{strip_prefix(top_jq[4]["名称"])}{top_jq[4]["数量"]}起（占比{top_jq[4]["占比"]}%）。城区两所（西门、东门派出所）共接报110起，占治安警情总量的56.1%，城区仍为治安防控重点区域。'
create_level3_paragraph(doc, '4.从辖区分布分析。', body_text)

# 小结
summary = (
    f'10月份治安警情{za["本期总量"]}起，环比上升{za["环比变化率"]}%，呈小幅反弹态势。'
    f'此类警情特征突出：'
)
p_summary = create_body_paragraph(doc, 'both')
set_first_line_indent(p_summary, 643, 200)
add_run_with_format(p_summary, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_summary, summary, '仿宋', 32, False, '000000')
add_run_with_format(p_summary, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary, f'盗窃和殴打他人两类案件高发，分别占治安警情的33.7%和25.5%，合计占比近六成；', '仿宋', 32, False, '000000')
add_run_with_format(p_summary, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary, f'凌晨和夜间时段发案集中，占比达34.6%，需加强夜间巡逻管控力度；', '仿宋', 32, False, '000000')
add_run_with_format(p_summary, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary, f'城区两所辖区仍为治安案件高发区域，占比56.1%，应持续强化城区重点区域治安巡防。建议各所队加大盗窃案件侦破力度和夜间巡逻密度，切实提升群众安全感。', '仿宋', 32, False, '000000')


# ---- (二) 治安殴打他人警情分析 ----
od = za['殴打他人分析']
create_level2_title(doc, '（二）治安殴打他人警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '治安殴打他人警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{od["本期数量"]}起，环比下降{abs(od["环比变化率"])}%。', '仿宋', 32, False, '000000')

# 1. 从发生原因分析
reasons = od['原因分布']
body_text = f'殴打他人警情主要原因为{reasons[0]["原因"]}{reasons[0]["数量"]}起（占比{reasons[0]["占比"]}%），其次{reasons[2]["原因"]}{reasons[2]["数量"]}起（占比{reasons[2]["占比"]}%），{reasons[3]["原因"]}{reasons[3]["数量"]}起（占比{reasons[3]["占比"]}%），{reasons[4]["原因"]}{reasons[4]["数量"]}起（占比{reasons[4]["占比"]}%），{reasons[5]["原因"]}{reasons[5]["数量"]}起（占比{reasons[5]["占比"]}%），{reasons[6]["原因"]}{reasons[6]["数量"]}起（占比{reasons[6]["占比"]}%）。口角琐事引发的殴打案件占比过半，反映出日常矛盾纠纷未能及时化解。'
create_level3_paragraph(doc, '1.从发生原因分析。', body_text)

# 2. 从涉刀警情分析
knife = od['涉刀警情']
body_text = f'殴打他人警情中涉刀警情本月{knife["本期"]}起，上月{knife["上期"]}起，环比上升{round((knife["本期"]-knife["上期"])/knife["上期"]*100, 1)}%。涉刀殴打案件占殴打他人警情的{round(knife["本期"]/od["本期数量"]*100, 1)}%，涉刀比例较高，潜在危害性大，需重点关注。'
create_level3_paragraph(doc, '2.从涉刀警情分析。', body_text)

# 3. 从发案时段分析
ot = od['时段分布']
ot_sorted = sorted(ot, key=lambda x: x['数量'], reverse=True)
body_text = f'殴打他人警情发案时段主要集中在{ot_sorted[0]["时段"]}{ot_sorted[0]["数量"]}起（占比{ot_sorted[0]["占比"]}%），其次{ot_sorted[1]["时段"]}{ot_sorted[1]["数量"]}起（占比{ot_sorted[1]["占比"]}%），{ot_sorted[2]["时段"]}{ot_sorted[2]["数量"]}起（占比{ot_sorted[2]["占比"]}%）。凌晨时段高发，占比达32%，与酒后闹事密切相关。'
create_level3_paragraph(doc, '3.从发案时段分析。', body_text)

# 4. 从辖区分布分析
od_jq = od['辖区分布']
top_od = od_jq[:5]
body_text = f'殴打他人警情主要集中在{strip_prefix(top_od[0]["名称"])}{top_od[0]["数量"]}起（占比{top_od[0]["占比"]}%），其次{strip_prefix(top_od[1]["名称"])}{top_od[1]["数量"]}起（占比{top_od[1]["占比"]}%），{strip_prefix(top_od[2]["名称"])}{top_od[2]["数量"]}起（占比{top_od[2]["占比"]}%），{strip_prefix(top_od[3]["名称"])}{top_od[3]["数量"]}起（占比{top_od[3]["占比"]}%），{strip_prefix(top_od[4]["名称"])}{top_od[4]["数量"]}起（占比{top_od[4]["占比"]}%）。'
create_level3_paragraph(doc, '4.从辖区分布分析。', body_text)

# 小结
p_summary2 = create_body_paragraph(doc, 'both')
set_first_line_indent(p_summary2, 643, 200)
add_run_with_format(p_summary2, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_summary2, f'10月份殴打他人警情{od["本期数量"]}起，虽环比小幅下降{abs(od["环比变化率"])}%，但仍处于较高水平。此类警情特征突出：', '仿宋', 32, False, '000000')
add_run_with_format(p_summary2, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary2, f'口角琐事为主要诱因，占比50%，日常矛盾排查化解工作仍需加强；', '仿宋', 32, False, '000000')
add_run_with_format(p_summary2, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary2, f'凌晨时段高发且涉刀比例达20%，安全隐患突出；', '仿宋', 32, False, '000000')
add_run_with_format(p_summary2, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_summary2, f'城区两所（西门、东门）占比54%，仍是殴打案件高发区域。建议加强城区夜间娱乐场所及重点区域巡逻管控，及时发现和制止暴力行为。', '仿宋', 32, False, '000000')


# ---- (三) 刑事警情分析 ----
xs = data['各类详细分析']['刑事警情']
create_level2_title(doc, '（三）刑事警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '刑事警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{xs["本期总量"]}起，环比上升{xs["环比变化率"]}%。', '仿宋', 32, False, '000000')

# 1. 从警情类型分析
xs_types = xs['反馈报警类型分布']
xs_sorted = sorted(xs_types, key=lambda x: x['本期数量'], reverse=True)
body_text = '其中'
for i, t in enumerate(xs_sorted):
    if t['本期数量'] == 0:
        continue
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["类型"]}{t["本期数量"]}起，{rate_text}'
    body_text += '；' if i < len([x for x in xs_sorted if x['本期数量'] > 0]) - 1 else '。'
create_level3_paragraph(doc, '1.从警情类型分析。', body_text)

# 2. 从高发案件类型分析
xs_fine = xs['反馈报警细类分布']
body_text = f'刑事警情中高发类型为{xs_fine[0]["名称"]}{xs_fine[0]["数量"]}起（占比{xs_fine[0]["占比"]}%），其次{xs_fine[1]["名称"]}{xs_fine[1]["数量"]}起（占比{xs_fine[1]["占比"]}%）。'

# 盗窃子类
theft = xs['盗窃分析']
theft_sub = theft['盗窃子类分布']
body_text += f'刑事盗窃案件中，{theft_sub[0]["名称"]}{theft_sub[0]["数量"]}起（占比{theft_sub[0]["占比"]}%）为最高发类型，其次{theft_sub[1]["名称"]}{theft_sub[1]["数量"]}起（占比{theft_sub[1]["占比"]}%）。'

# 电诈
epz = xs['电信网络诈骗分析']
epz_sub = epz['诈骗子类分布']
body_text += f'电信网络诈骗{epz["本期数量"]}起，环比上升{epz["环比变化率"]}%，其中{epz_sub[0]["名称"]}{epz_sub[0]["数量"]}起（占比{epz_sub[0]["占比"]}%）。'
create_level3_paragraph(doc, '2.从高发案件类型分析。', body_text)

# 3. 从辖区分布分析
xs_jq = xs['辖区分布']
top_xs = xs_jq[:5]
body_text = f'刑事警情主要分布在{strip_prefix(top_xs[0]["名称"])}{top_xs[0]["数量"]}起（占比{top_xs[0]["占比"]}%），其次{strip_prefix(top_xs[1]["名称"])}{top_xs[1]["数量"]}起（占比{top_xs[1]["占比"]}%），{strip_prefix(top_xs[2]["名称"])}{top_xs[2]["数量"]}起（占比{top_xs[2]["占比"]}%），{strip_prefix(top_xs[3]["名称"])}{top_xs[3]["数量"]}起（占比{top_xs[3]["占比"]}%），{strip_prefix(top_xs[4]["名称"])}{top_xs[4]["数量"]}起（占比{top_xs[4]["占比"]}%）。'
create_level3_paragraph(doc, '3.从辖区分布分析。', body_text)

# 小结
p_xs_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_xs_sum, 643, 200)
add_run_with_format(p_xs_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_xs_sum, f'10月份刑事警情{xs["本期总量"]}起，环比小幅上升{xs["环比变化率"]}%。此类警情特征突出：', '仿宋', 32, False, '000000')
add_run_with_format(p_xs_sum, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_xs_sum, f'侵财类案件占比突出，盗窃和电信网络诈骗合计占刑事警情的78.9%；', '仿宋', 32, False, '000000')
add_run_with_format(p_xs_sum, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_xs_sum, f'入户盗窃占刑事盗窃的50%，反映出居民住宅安防仍存薄弱环节；', '仿宋', 32, False, '000000')
add_run_with_format(p_xs_sum, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_xs_sum, f'电信网络诈骗环比上升16.7%，反诈宣传教育工作仍需持续推进。建议各所队加大入户盗窃案件侦破力度，强化社区安防宣传，持续开展反诈宣传进村入户。', '仿宋', 32, False, '000000')


# ---- (四) 交通警情分析 ----
jt = data['各类详细分析']['交通警情']
create_level2_title(doc, '（四）交通警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '交通警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{jt["本期总量"]}起，环比上升{jt["环比变化率"]}%，上升幅度较大。', '仿宋', 32, False, '000000')

# 1. 从警情类型分析
jt_types = jt['反馈报警类型分布']
jt_sorted = sorted(jt_types, key=lambda x: x['本期数量'], reverse=True)
body_text = '其中'
for i, t in enumerate(jt_sorted):
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["类型"]}{t["本期数量"]}起，{rate_text}'
    if i < len(jt_sorted) - 1:
        body_text += '；'
    else:
        body_text += '。'
body_text += f'道路交通事故仍为交通警情的绝对主体，占比{jt["反馈报警细类分布"][0]["占比"]}%。交通事故逃逸{jt["交通事故逃逸"]["本期"]}起，环比上升{jt["交通事故逃逸"]["环比变化率"]}%。'
create_level3_paragraph(doc, '1.从警情类型分析。', body_text)

# 2. 从发案时段分析
jt_ts = jt['时段分布']
jt_ts_sorted = sorted(jt_ts, key=lambda x: x['数量'], reverse=True)
body_text = f'交通警情发案时段主要集中在{jt_ts_sorted[0]["时段"]}{jt_ts_sorted[0]["数量"]}起（占比{jt_ts_sorted[0]["占比"]}%），其次{jt_ts_sorted[1]["时段"]}{jt_ts_sorted[1]["数量"]}起（占比{jt_ts_sorted[1]["占比"]}%），{jt_ts_sorted[2]["时段"]}{jt_ts_sorted[2]["数量"]}起（占比{jt_ts_sorted[2]["占比"]}%）。下午至晚高峰时段（14:00-20:59）为事故高峰期，合计{jt_ts[4]["数量"]+jt_ts[5]["数量"]}起，占比{round((jt_ts[4]["数量"]+jt_ts[5]["数量"])/jt["本期总量"]*100,1)}%。'
create_level3_paragraph(doc, '2.从发案时段分析。', body_text)

# 3. 从周末工作日分析
wd = jt['周末工作日分布']
body_text = f'交通警情中工作日发案{wd["工作日数量"]}起，占比{wd["工作日占比"]}%；周末发案{wd["周末数量"]}起，占比{wd["周末占比"]}%。10月份含国庆长假，假期出行增加导致交通事故显著上升。'
create_level3_paragraph(doc, '3.从工作日与周末分析。', body_text)

# 4. 从高发细类分析
jt_fine = jt['反馈报警细类分布']
body_text = f'交通警情中高发细类为{jt_fine[0]["名称"]}{jt_fine[0]["数量"]}起（占比{jt_fine[0]["占比"]}%），其次{jt_fine[1]["名称"]}{jt_fine[1]["数量"]}起（占比{jt_fine[1]["占比"]}%），{jt_fine[2]["名称"]}{jt_fine[2]["数量"]}起（占比{jt_fine[2]["占比"]}%），{jt_fine[3]["名称"]}{jt_fine[3]["数量"]}起（占比{jt_fine[3]["占比"]}%），{jt_fine[4]["名称"]}{jt_fine[4]["数量"]}起（占比{jt_fine[4]["占比"]}%）。交通违法行为{jt["交通违法"]["本期"]}起，环比大幅上升{jt["交通违法"]["环比变化率"]}%，交通秩序管理压力增大。'
create_level3_paragraph(doc, '4.从高发细类分析。', body_text)

# 小结
p_jt_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_jt_sum, 643, 200)
add_run_with_format(p_jt_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_jt_sum, f'10月份交通警情{jt["本期总量"]}起，环比大幅上升{jt["环比变化率"]}%，为本月上升幅度最大的警情类别。此类警情特征突出：', '仿宋', 32, False, '000000')
add_run_with_format(p_jt_sum, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_jt_sum, f'国庆长假期间出行量激增，道路交通事故大幅增加，环比上升32.2%；', '仿宋', 32, False, '000000')
add_run_with_format(p_jt_sum, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_jt_sum, f'交通违法行为环比上升93.8%，违法停车和违法通行问题突出；', '仿宋', 32, False, '000000')
add_run_with_format(p_jt_sum, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_jt_sum, f'下午至晚高峰时段为事故高发期，占比46.9%。建议交通管理大队持续加强重点路段和时段的执法管控，加大对违法停车、违法通行等交通违法行为的查处力度，提升道路交通安全管理水平。', '仿宋', 32, False, '000000')


# ---- (五) 纠纷警情分析 ----
jf = data['各类详细分析']['纠纷警情']
create_level2_title(doc, '（五）纠纷警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '纠纷警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{jf["本期总量"]}起，环比上升{jf["环比变化率"]}%。', '仿宋', 32, False, '000000')

# 1. 从纠纷类型分析
jf_types = jf['反馈报警类型分布']
jf_sorted = sorted(jf_types, key=lambda x: x['本期数量'], reverse=True)
body_text = '其中'
for i, t in enumerate(jf_sorted[:6]):
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["类型"]}{t["本期数量"]}起，{rate_text}'
    if i < 5:
        body_text += '；'
    else:
        body_text += '。'
create_level3_paragraph(doc, '1.从纠纷类型分析。', body_text)

# 2. 从高发纠纷细类分析
jf_fine = jf['反馈报警细类分布']
top_jf_fine = jf_fine[:6]
body_text = f'纠纷警情中高发细类主要为{top_jf_fine[0]["名称"]}{top_jf_fine[0]["数量"]}起（占比{top_jf_fine[0]["占比"]}%），其次{top_jf_fine[1]["名称"]}{top_jf_fine[1]["数量"]}起（占比{top_jf_fine[1]["占比"]}%），{top_jf_fine[2]["名称"]}{top_jf_fine[2]["数量"]}起（占比{top_jf_fine[2]["占比"]}%），{top_jf_fine[3]["名称"]}{top_jf_fine[3]["数量"]}起（占比{top_jf_fine[3]["占比"]}%），{top_jf_fine[4]["名称"]}{top_jf_fine[4]["数量"]}起（占比{top_jf_fine[4]["占比"]}%），{top_jf_fine[5]["名称"]}{top_jf_fine[5]["数量"]}起（占比{top_jf_fine[5]["占比"]}%）。土地权属和经济类纠纷占比较高，涉及群众切身利益，需重点关注。'
create_level3_paragraph(doc, '2.从高发纠纷细类分析。', body_text)

# 3. 从发案时段分析
jf_ts = jf['时段分布']
jf_ts_sorted = sorted(jf_ts, key=lambda x: x['数量'], reverse=True)
body_text = f'纠纷警情发案时段主要集中在{jf_ts_sorted[0]["时段"]}{jf_ts_sorted[0]["数量"]}起（占比{jf_ts_sorted[0]["占比"]}%），其次{jf_ts_sorted[1]["时段"]}{jf_ts_sorted[1]["数量"]}起（占比{jf_ts_sorted[1]["占比"]}%），{jf_ts_sorted[2]["时段"]}{jf_ts_sorted[2]["数量"]}起（占比{jf_ts_sorted[2]["占比"]}%）。白天时段（09:00-17:59）为纠纷高发期，合计{jf_ts[2]["数量"]+jf_ts[3]["数量"]+jf_ts[4]["数量"]}起，占比{round((jf_ts[2]["数量"]+jf_ts[3]["数量"]+jf_ts[4]["数量"])/jf["本期总量"]*100,1)}%。'
create_level3_paragraph(doc, '3.从发案时段分析。', body_text)

# 4. 从辖区分布分析
jf_jq = jf['辖区分布']
top_jf_jq = jf_jq[:6]
body_text = f'纠纷警情主要集中在{strip_prefix(top_jf_jq[0]["名称"])}{top_jf_jq[0]["数量"]}起（占比{top_jf_jq[0]["占比"]}%），其次{strip_prefix(top_jf_jq[1]["名称"])}{top_jf_jq[1]["数量"]}起（占比{top_jf_jq[1]["占比"]}%），{strip_prefix(top_jf_jq[2]["名称"])}{top_jf_jq[2]["数量"]}起（占比{top_jf_jq[2]["占比"]}%），{strip_prefix(top_jf_jq[3]["名称"])}{top_jf_jq[3]["数量"]}起（占比{top_jf_jq[3]["占比"]}%），{strip_prefix(top_jf_jq[4]["名称"])}{top_jf_jq[4]["数量"]}起（占比{top_jf_jq[4]["占比"]}%），{strip_prefix(top_jf_jq[5]["名称"])}{top_jf_jq[5]["数量"]}起（占比{top_jf_jq[5]["占比"]}%）。'
create_level3_paragraph(doc, '4.从辖区分布分析。', body_text)

# 小结
p_jf_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_jf_sum, 643, 200)
add_run_with_format(p_jf_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_jf_sum, f'10月份纠纷警情{jf["本期总量"]}起，环比上升{jf["环比变化率"]}%。此类警情特征突出：', '仿宋', 32, False, '000000')
add_run_with_format(p_jf_sum, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_jf_sum, f'其他纠纷和邻里纠纷环比涨幅较大，分别上升109.5%和84.6%，需重点关注基层矛盾化解工作；', '仿宋', 32, False, '000000')
add_run_with_format(p_jf_sum, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_jf_sum, f'土地权属纠纷和拖欠工资问题仍较突出，涉及群众切身利益；', '仿宋', 32, False, '000000')
add_run_with_format(p_jf_sum, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_jf_sum, f'城区两所辖区纠纷警情占比41%，基层社区矛盾调处机制需进一步完善。建议各所队深入推进矛盾纠纷排查化解工作，对土地权属、劳资等重点领域纠纷加强源头治理。', '仿宋', 32, False, '000000')


# ---- (六) 群众紧急求助分析 ----
qz = data['各类详细分析']['群众紧急求助']
create_level2_title(doc, '（六）群众紧急求助分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '群众紧急求助', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{qz["本期总量"]}起，环比上升{qz["环比变化率"]}%。', '仿宋', 32, False, '000000')

# 1. 从求助类型分析
qz_types = qz['反馈报警类型分布']
qz_sorted = sorted(qz_types, key=lambda x: x['本期数量'], reverse=True)
body_text = '其中'
for i, t in enumerate(qz_sorted[:5]):
    if t['本期数量'] == 0:
        continue
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["类型"]}{t["本期数量"]}起，{rate_text}'
    body_text += '；'
body_text = body_text.rstrip('；') + '。'
body_text += f'其他紧急求助占群众紧急求助总量的{round(542/qz["本期总量"]*100,1)}%，为绝对主体。挪车求助{164}起，环比大幅上升76.3%，与国庆假期停车位紧张密切相关。'
create_level3_paragraph(doc, '1.从求助类型分析。', body_text)

# 2. 从辖区分布分析
qz_jq = qz['辖区分布']
top_qz = qz_jq[:5]
body_text = f'群众紧急求助主要集中在{strip_prefix(top_qz[0]["名称"])}{top_qz[0]["数量"]}起（占比{top_qz[0]["占比"]}%），其次{strip_prefix(top_qz[1]["名称"])}{top_qz[1]["数量"]}起（占比{top_qz[1]["占比"]}%），{strip_prefix(top_qz[2]["名称"])}{top_qz[2]["数量"]}起（占比{top_qz[2]["占比"]}%），{strip_prefix(top_qz[3]["名称"])}{top_qz[3]["数量"]}起（占比{top_qz[3]["占比"]}%），{strip_prefix(top_qz[4]["名称"])}{top_qz[4]["数量"]}起（占比{top_qz[4]["占比"]}%）。城区两所合计351起，占比59.5%。'
create_level3_paragraph(doc, '2.从辖区分布分析。', body_text)

# 小结
p_qz_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_qz_sum, 643, 200)
add_run_with_format(p_qz_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_qz_sum, f'10月份群众紧急求助{qz["本期总量"]}起，环比上升{qz["环比变化率"]}%。此类警情特征突出：', '仿宋', 32, False, '000000')
add_run_with_format(p_qz_sum, '一是', '仿宋', 32, True, '000000')
add_run_with_format(p_qz_sum, f'挪车求助环比大幅上升76.3%，节假日期间停车矛盾加剧；', '仿宋', 32, False, '000000')
add_run_with_format(p_qz_sum, '二是', '仿宋', 32, True, '000000')
add_run_with_format(p_qz_sum, f'贵重失物求助环比上升57.6%，节假日人员流动增加导致遗失物品情况增多；', '仿宋', 32, False, '000000')
add_run_with_format(p_qz_sum, '三是', '仿宋', 32, True, '000000')
add_run_with_format(p_qz_sum, f'城区两所占比近六成，城区公共服务需求持续旺盛。建议加强城区公共停车设施建设和管理，优化挪车求助处理机制。', '仿宋', 32, False, '000000')


# ---- (七) 其他警情分析 ----
qt = data['各类详细分析']['其他警情']
create_level2_title(doc, '（七）其他警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份我局共接报', '仿宋', 32, False, '000000')
add_run_with_format(p, '其他警情', '仿宋', 32, True, '000000')
add_run_with_format(p, f'{qt["本期总量"]}起，环比上升{qt["环比变化率"]}%。', '仿宋', 32, False, '000000')

# 1. 从子类别分析
qt_sub = qt['子类别分布']
qt_sorted = sorted(qt_sub, key=lambda x: x['本期数量'], reverse=True)
body_text = '其中'
for i, t in enumerate(qt_sorted):
    if t['本期数量'] == 0:
        continue
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["子类别"]}{t["本期数量"]}起，{rate_text}'
    body_text += '；'
body_text = body_text.rstrip('；') + '。'
create_level3_paragraph(doc, '1.从子类别分析。', body_text)

# 小结
p_qt_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_qt_sum, 643, 200)
add_run_with_format(p_qt_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_qt_sum, f'10月份其他警情{qt["本期总量"]}起，环比小幅上升{qt["环比变化率"]}%，整体保持平稳。咨询类和举报类警情环比有所下降，社会联动类出现4起新增。', '仿宋', 32, False, '000000')


# ---- (八) 金牌港重点园区分析 ----
jp = data['金牌港重点园区']
create_level2_title(doc, '（八）金牌港重点园区警情分析')

# 总量概述
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份', '仿宋', 32, False, '000000')
add_run_with_format(p, '金牌港重点园区', '仿宋', 32, True, '000000')
add_run_with_format(p, f'共接报警情{jp["本期总量"]}起，环比下降{abs(jp["环比变化率"])}%。', '仿宋', 32, False, '000000')

# 1. 从警情类型分析
jp_types = jp['按大类分布']
body_text = '其中'
for i, t in enumerate(jp_types):
    rate_text = f'环比上升{t["环比变化率"]}%' if t["环比变化率"] > 0 else (
        f'环比下降{abs(t["环比变化率"])}%' if t["环比变化率"] < 0 else '与上月持平')
    body_text += f'{t["类别"]}{t["本期数量"]}起，{rate_text}'
    if i < len(jp_types) - 1:
        body_text += '；'
    else:
        body_text += '。'
body_text += f'交通警情仍为园区警情主体，占比{round(17/jp["本期总量"]*100, 1)}%。'
create_level3_paragraph(doc, '1.从警情类型分析。', body_text)

# 2. 从高发细类分析
jp_fine = jp['反馈报警类型分布_本期']
body_text = f'园区警情高发细类为{jp_fine[0]["名称"]}{jp_fine[0]["数量"]}起（占比{jp_fine[0]["占比"]}%），其次{jp_fine[1]["名称"]}{jp_fine[1]["数量"]}起（占比{jp_fine[1]["占比"]}%），{jp_fine[2]["名称"]}{jp_fine[2]["数量"]}起（占比{jp_fine[2]["占比"]}%），{jp_fine[3]["名称"]}{jp_fine[3]["数量"]}起（占比{jp_fine[3]["占比"]}%）。'
create_level3_paragraph(doc, '2.从高发细类分析。', body_text)

# 小结
p_jp_sum = create_body_paragraph(doc, 'both')
set_first_line_indent(p_jp_sum, 643, 200)
add_run_with_format(p_jp_sum, '小结：', '仿宋', 32, True, '000000')
add_run_with_format(p_jp_sum, f'10月份金牌港重点园区警情{jp["本期总量"]}起，环比下降{abs(jp["环比变化率"])}%，管控成效显著。交通事故仍为园区主要警情类型，建议持续加强园区道路交通安全管理和企业安全生产监督。', '仿宋', 32, False, '000000')


# =========================
# 第三章：下降警情类型分布
# =========================
create_level1_title(doc, '三、下降警情类型分布')

# 本月6大类警情环比均为上升，无下降类别
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '10月份六大类警情环比均呈上升态势，无环比下降的警情类别。', '仿宋', 32, False, '000000')


# =========================
# 第四章：工作建议
# =========================
create_level1_title(doc, '四、工作建议')

# (一) 强化城区治安巡防力度
create_level2_title(doc, '（一）强化城区治安巡防力度。')
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '临城西门、东门派出所辖区治安警情和殴打他人警情占比均超过50%，城区治安压力较大。建议两所加大辖区内重点区域、重点时段（凌晨及夜间）的巡逻防控力度，提高见警率和管事率。对辖区内的娱乐场所、夜市、出租屋等治安复杂场所，开展定期清查和隐患排查，及时发现和消除治安隐患。同时，强化社区民警入户走访工作，掌握辖区治安动态，将矛盾纠纷化解在萌芽状态。', '仿宋', 32, False, '000000')

# (二) 加大盗窃案件打击力度
create_level2_title(doc, '（二）加大盗窃案件打击力度。')
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '治安盗窃警情66起，环比上升17.9%，其中盗窃电动车22起占比33.3%为最高发类型，刑事盗窃中入户盗窃4起占比50%。建议各所队针对电动车盗窃高发问题，加强重点区域视频巡查和蹲点守候，对停放集中区域增设监控设备。对入户盗窃案件，组织力量快侦快办，加大打击力度形成震慑。深入开展安防宣传，引导群众加强电动车防盗锁具使用和居家安全防范意识。', '仿宋', 32, False, '000000')

# (三) 持续推进道路交通安全综合治理
create_level2_title(doc, '（三）持续推进道路交通安全综合治理。')
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '交通警情882起，环比大幅上升35.1%，为本月上升幅度最大的警情类别，交通违法行为环比上升93.8%。建议交通管理大队加强对重点路段、事故多发路段的日常巡逻和隐患排查，在下午至晚高峰时段（14:00-20:59）增加警力部署。加大对违法停车、无证驾驶、酒后驾驶等交通违法行为的查处力度，常态化开展交通安全宣传教育活动，切实提升辖区道路交通安全水平。节假日期间提前制定交通疏导预案，确保道路安全畅通。', '仿宋', 32, False, '000000')

# (四) 深化矛盾纠纷排查化解
create_level2_title(doc, '（四）深化矛盾纠纷排查化解。')
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '纠纷警情224起，环比上升13.1%，其中其他纠纷环比上升109.5%、邻里纠纷环比上升84.6%，增幅较大。建议各所队全面加强矛盾纠纷排查和源头治理工作，对土地权属、拖欠工资等群众反映强烈的问题，联合相关职能部门协调处置。健全完善多元化解机制，充分发挥人民调解组织作用，推进警调对接、访调对接，努力将矛盾化解在基层、消除在萌芽。对家庭婚姻情感纠纷等涉及个人极端行为风险的案件，建立跟踪回访机制，防止矛盾升级。', '仿宋', 32, False, '000000')

# (五) 持续强化反诈宣传和涉刀警情管控
create_level2_title(doc, '（五）持续强化反诈宣传和涉刀警情管控。')
p = create_body_paragraph(doc, 'both')
add_run_with_format(p, '电信网络诈骗警情（刑事+治安）共17起，环比上升30.8%；涉刀警情61起，虽环比下降7.6%，但基数仍较大。建议各所队持续深入开展反诈宣传进村入户活动，针对冒充电商客服、刷单返利等高发诈骗类型，加强精准宣防。同时，加大对涉刀警情的重视程度，对殴打他人涉刀案件（10起，占殴打警情的20%）高度关注，强化刀具管控和重点人员管理，对扬言实施极端行为等敏感警情做到快速响应、妥善处置。', '仿宋', 32, False, '000000')


# ============================================================
# 落款区域
# ============================================================

# 空行
for _ in range(2):
    p_blank = doc.add_paragraph()
    set_line_spacing_exact(p_blank, 560)

# 落款单位名
p_sign = doc.add_paragraph()
set_line_spacing_exact(p_sign, 560)
set_left_indent(p_sign, 4160)
add_run_with_format(p_sign, '临高县公安局情报指挥中心', '仿宋', 32, False, '000000')

# 落款日期
p_date = doc.add_paragraph()
set_line_spacing_exact(p_date, 560)
set_left_indent(p_date, 4800)
add_run_with_format(p_date, '2025年10月31日', '仿宋', 32, False, '000000')

# 空行用于分页留白
for _ in range(4):
    p_blank = doc.add_paragraph()
    set_line_spacing_exact(p_blank, 560)

# ============================================================
# 页尾区域：抄送/抄报/印发
# ============================================================

# 上方黑色分隔线
p_line1 = add_black_line_top(doc)
add_run_with_format(p_line1, '抄报：县局领导。', '仿宋', 28, False, '000000')

# 抄送行 + 下方分隔线
p_line2 = add_black_line_bottom(doc)
add_run_with_format(p_line2, '抄送：各所、队、室（中心）。', '仿宋', 28, False, '000000')

# 印发行
p_print = doc.add_paragraph()
set_line_spacing_exact(p_print, 560)
add_run_with_format(p_print, '临高县公安局情报指挥中心', '仿宋', 28, False, '000000')
# 右侧日期
add_run_with_format(p_print, '                    ', '仿宋', 28, False, '000000')
add_run_with_format(p_print, '2025年10月31日印发', '仿宋', 28, False, '000000')


# ============================================================
# 保存文档
# ============================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"报告已生成: {OUTPUT_FILE}")
print(f"文件大小: {os.path.getsize(OUTPUT_FILE)} bytes")
