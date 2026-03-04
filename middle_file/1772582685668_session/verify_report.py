#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证生成的报告格式和内容
"""

from docx import Document
from docx.oxml.ns import qn
import os

OUTPUT_FILE = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_2025年10月统计报告.docx"

doc = Document(OUTPUT_FILE)

print("=" * 60)
print("报告验证报告")
print("=" * 60)

# 1. 基本信息
print(f"\n1. 基本信息")
print(f"   文件大小: {os.path.getsize(OUTPUT_FILE)} bytes")
print(f"   段落总数: {len(doc.paragraphs)}")
print(f"   页面宽度: {doc.sections[0].page_width}")
print(f"   页面高度: {doc.sections[0].page_height}")

# 2. 检查关键段落内容
print(f"\n2. 关键段落内容检查")
for i, p in enumerate(doc.paragraphs[:5]):
    text = p.text.strip()
    if text:
        print(f"   段落{i}: {text[:60]}...")

# 3. 检查行距设置
print(f"\n3. 行距设置检查 (抽样前20个段落)")
line_issues = 0
for i, p in enumerate(doc.paragraphs[:20]):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line_val = spacing.get(qn('w:line'))
            line_rule = spacing.get(qn('w:lineRule'))
            if line_val and int(line_val) > 1000:
                print(f"   !! 段落{i} 行距异常: w:line={line_val} (可能误用了EMU)")
                line_issues += 1
            elif line_val:
                pass  # 正常
if line_issues == 0:
    print(f"   OK - 所有抽样段落行距正常 (值在合理twips范围内)")
else:
    print(f"   WARNING - 发现{line_issues}个行距异常段落!")

# 4. 检查字号设置
print(f"\n4. 字号设置检查")
font_sizes = set()
for p in doc.paragraphs:
    for run in p.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                font_sizes.add(val)
print(f"   检测到的字号半磅值: {sorted(font_sizes)}")
expected = {'110', '44', '32', '28'}
found = font_sizes & expected
missing = expected - font_sizes
print(f"   期望字号: {sorted(expected)}")
print(f"   已找到: {sorted(found)}")
if missing:
    print(f"   缺失: {sorted(missing)}")
else:
    print(f"   OK - 所有期望字号均已使用")

# 5. 检查字体
print(f"\n5. 字体设置检查")
fonts = set()
for p in doc.paragraphs:
    for run in p.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    fonts.add(ea)
print(f"   检测到的字体: {sorted(fonts)}")
expected_fonts = {'方正小标宋简体', '黑体', '楷体', '仿宋'}
found_fonts = fonts & expected_fonts
missing_fonts = expected_fonts - fonts
print(f"   期望字体: {sorted(expected_fonts)}")
print(f"   已找到: {sorted(found_fonts)}")
if missing_fonts:
    print(f"   缺失: {sorted(missing_fonts)}")
else:
    print(f"   OK - 所有期望字体均已使用")

# 6. 检查章节结构
print(f"\n6. 章节结构检查")
chapter_markers = ['一、', '二、', '三、', '四、']
found_chapters = []
for p in doc.paragraphs:
    text = p.text.strip()
    for marker in chapter_markers:
        if text.startswith(marker):
            found_chapters.append(text[:20])
print(f"   找到一级标题: {found_chapters}")
if len(found_chapters) == 4:
    print(f"   OK - 四个一级标题完整")
else:
    print(f"   WARNING - 期望4个一级标题, 实际{len(found_chapters)}个")

# 7. 检查二级标题
print(f"\n7. 二级标题检查")
sub_chapters = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('（') and '）' in text[:5]:
        sub_chapters.append(text[:30])
print(f"   找到二级标题数: {len(sub_chapters)}")
for sc in sub_chapters:
    print(f"     - {sc}")

# 8. 检查小结
print(f"\n8. 小结段落检查")
summaries = 0
for p in doc.paragraphs:
    if '小结：' in p.text:
        summaries += 1
print(f"   找到小结段落数: {summaries}")

# 9. 检查加粗关键词
print(f"\n9. 加粗关键词检查 (抽样)")
bold_keywords = []
for p in doc.paragraphs[:30]:
    for run in p.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            b = rPr.find(qn('w:b'))
            if b is not None and run.text.strip():
                bold_keywords.append(run.text.strip()[:20])
print(f"   前30段中加粗的关键词数: {len(bold_keywords)}")
if bold_keywords[:10]:
    for kw in bold_keywords[:10]:
        print(f"     - {kw}")

# 10. 整体评估
print(f"\n{'=' * 60}")
print("整体评估")
print(f"{'=' * 60}")
issues = 0
if line_issues > 0:
    issues += 1
    print("  [FAIL] 行距设置存在问题")
else:
    print("  [PASS] 行距设置正确 (使用twips原始值)")

if missing:
    issues += 1
    print("  [FAIL] 字号设置不完整")
else:
    print("  [PASS] 字号设置完整")

if missing_fonts:
    issues += 1
    print("  [FAIL] 字体设置不完整")
else:
    print("  [PASS] 字体设置完整")

if len(found_chapters) != 4:
    issues += 1
    print("  [FAIL] 章节结构不完整")
else:
    print("  [PASS] 四大章节完整")

if summaries < 5:
    issues += 1
    print(f"  [WARN] 小结段落较少 ({summaries}个)")
else:
    print(f"  [PASS] 小结段落充足 ({summaries}个)")

if len(bold_keywords) < 5:
    issues += 1
    print("  [WARN] 加粗关键词较少")
else:
    print("  [PASS] 加粗关键词设置正常")

if issues == 0:
    print(f"\n  总评: 全部通过!")
else:
    print(f"\n  总评: 发现{issues}个问题需要关注")
