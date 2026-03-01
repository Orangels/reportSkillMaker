from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# 读取数据
with open('./middle_file/oct2025/extracted_data_oct_corrected.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(16)

# 1. 发文单位（红色，居中，方正小标宋简体 55pt）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('临高县公安局情报指挥中心')
run.font.size = Pt(55)
run.font.color.rgb = RGBColor(255, 0, 0)
run.font.bold = True

# 2. 主标题（居中，方正小标宋简体 22pt）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('关于10月份警情分析的报告')
run.font.size = Pt(22)
run.font.bold = True

# 3. 一、整体情况
p = doc.add_paragraph()
run = p.add_run('一、整体情况')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True

# 整体情况内容
curr = data['current_period']
prev = data['previous_period']
comp = data['comparison']

content = f"10月1日至30日我局共接报有效警情{curr['valid']}起（不含骚扰警情{curr['harassment']}起），环比上升{comp['valid']}%。"
content += f"其中刑事警情{curr['six_categories']['criminal']}起，环比上升{comp['criminal']}%；"
content += f"治安警情{curr['six_categories']['security']}起，环比上升{comp['security']}%；"
content += f"交通警情{curr['six_categories']['traffic']}起，环比上升{comp['traffic']}%；"
content += f"纠纷警情{curr['six_categories']['dispute']}起，环比上升{comp['dispute']}%；"
content += f"群众紧急求助{curr['six_categories']['help']}起，环比上升{comp['help']}%；"
content += f"其他警情{curr['six_categories']['other']}起，环比下降{abs(comp['other'])}%。"

p = doc.add_paragraph(content)
p.paragraph_format.first_line_indent = Inches(0.5)

# 4. 二、上升警情类别分布
p = doc.add_paragraph()
run = p.add_run('二、上升警情类别分布')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True

# 交通警情分析
p = doc.add_paragraph()
run = p.add_run('（一）交通警情分析')
run.font.name = '黑体'
run.font.size = Pt(16)

content = f"我局共接报交通警情{curr['six_categories']['traffic']}起，环比上升{comp['traffic']}%，上升幅度较大。"
p = doc.add_paragraph(content)
p.paragraph_format.first_line_indent = Inches(0.5)

# 小结
p = doc.add_paragraph()
run = p.add_run('小结：')
run.font.bold = True
run = p.add_run('近期交通警情环比大幅上升，需加强交通管控和安全宣传，持续压降交通警情。')
p.paragraph_format.first_line_indent = Inches(0.5)

# 5. 三、下降警情类型分布
p = doc.add_paragraph()
run = p.add_run('三、下降警情类型分布')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True

# 其他警情分析
p = doc.add_paragraph()
run = p.add_run('（一）其他警情分析')
run.font.name = '黑体'
run.font.size = Pt(16)

content = f"我局共接报其他警情{curr['six_categories']['other']}起，环比下降{abs(comp['other'])}%。"
p = doc.add_paragraph(content)
p.paragraph_format.first_line_indent = Inches(0.5)

# 6. 四、工作建议
p = doc.add_paragraph()
run = p.add_run('四、工作建议')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True

# 建议1
p = doc.add_paragraph()
run = p.add_run('（一）强化交通管控。')
run.font.name = '黑体'
run.font.size = Pt(16)

content = "交警部门要加强重点时段和路段的交通管控，强化交通安全宣传引导，持续压降交通警情。"
p = doc.add_paragraph(content)
p.paragraph_format.first_line_indent = Inches(0.5)

# 7. 落款
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('临高县公安局情报指挥中心')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2025年11月8日')
run.font.size = Pt(16)

# 8. 抄送抄报
p = doc.add_paragraph()
run = p.add_run('抄送：各所、队、室（中心）')
p.paragraph_format.first_line_indent = Inches(0)

p = doc.add_paragraph()
run = p.add_run('抄报：严树勋副县长，各局领导')

p = doc.add_paragraph()
run = p.add_run('临高县公安局情报指挥中心 2025年11月8日印发')

# 保存文档
doc.save('./output/output_2025年10月统计报告.docx')
print("报告生成完成！")
print("文件路径：./output/output_2025年10月统计报告.docx")
