#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证生成的报告：检查格式、内容、数据准确性
"""

import os
from docx import Document
from docx.oxml.ns import qn

OUTPUT_FILE = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_2025年12月统计报告.docx"

doc = Document(OUTPUT_FILE)

print("=" * 60)
print("验证报告格式和内容")
print("=" * 60)

# 1. 检查页面设置
section = doc.sections[0]
print(f"\n[页面设置]")
print(f"  页面宽度: {section.page_width} EMU")
print(f"  页面高度: {section.page_height} EMU")
print(f"  上边距: {section.top_margin} EMU")
print(f"  下边距: {section.bottom_margin} EMU")
print(f"  左边距: {section.left_margin} EMU")
print(f"  右边距: {section.right_margin} EMU")

# 2. 检查段落数和行距
print(f"\n[段落统计]")
print(f"  总段落数: {len(doc.paragraphs)}")

# 检查行距设置
line_spacing_errors = []
for i, para in enumerate(doc.paragraphs):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line_val = spacing.get(qn('w:line'))
            if line_val and int(line_val) > 1000:
                line_spacing_errors.append((i, line_val, para.text[:50]))

if line_spacing_errors:
    print(f"\n  [错误] 发现行距值异常的段落（可能使用了EMU而非twips）:")
    for idx, val, text in line_spacing_errors:
        print(f"    段落{idx}: w:line={val}, 内容: {text}")
else:
    print(f"  [通过] 所有段落行距值正常")

# 3. 检查字体
print(f"\n[字体检查]")
font_summary = {}
for para in doc.paragraphs:
    for run in para.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'), 'N/A')
                sz_elem = rPr.find(qn('w:sz'))
                sz = sz_elem.get(qn('w:val')) if sz_elem is not None else 'N/A'
                key = f"{ea} (sz={sz})"
                if key not in font_summary:
                    font_summary[key] = 0
                font_summary[key] += 1

for font, count in sorted(font_summary.items()):
    print(f"  {font}: {count}个run")

# 4. 检查内容结构
print(f"\n[内容结构检查]")
key_sections = [
    "临高县公安局情报指挥中心",
    "关于12月份",
    "一、整体情况",
    "二、上升警情类别分布",
    "（一）交通警情分析",
    "（二）纠纷警情分析",
    "三、下降警情类型分布",
    "（一）治安警情分析",
    "（二）刑事警情分析",
    "四、工作建议",
    "小结：",
    "抄送：",
    "抄报：",
]

for section_name in key_sections:
    found = False
    for para in doc.paragraphs:
        if section_name in para.text:
            found = True
            break
    status = "[通过]" if found else "[缺失]"
    print(f"  {status} {section_name}")

# 5. 检查数据准确性
print(f"\n[数据验证]")
# 检查关键数据是否出现
data_checks = [
    ("有效警情3013起", "有效警情总量"),
    ("骚扰警情1起", "骚扰警情数量"),
    ("刑事警情18起", "刑事警情数量"),
    ("治安警情157起", "治安警情数量"),
    ("交通警情923起", "交通警情数量"),
    ("纠纷警情281起", "纠纷警情数量"),
    ("群众紧急求助858起", "群众紧急求助数量"),
    ("其他警情777起", "其他警情数量"),
]

full_text = "\n".join([p.text for p in doc.paragraphs])

for data_str, desc in data_checks:
    found = data_str in full_text
    status = "[通过]" if found else "[缺失]"
    print(f"  {status} {desc}: {data_str}")

# 6. 检查加粗标记
print(f"\n[加粗检查]")
bold_keywords = ["有效警情", "交通警情", "纠纷警情", "治安警情", "刑事警情", "小结：", "一是", "二是", "三是"]
for keyword in bold_keywords:
    found_bold = False
    for para in doc.paragraphs:
        for run in para.runs:
            if keyword in run.text and run.font.bold:
                found_bold = True
                break
        if found_bold:
            break
    status = "[通过]" if found_bold else "[缺失]"
    print(f"  {status} '{keyword}' 加粗")

# 7. 统计段落详情
print(f"\n[段落内容概览（前30段）]")
for i, para in enumerate(doc.paragraphs[:30]):
    text_preview = para.text[:80] if para.text else "(空段落)"
    align = str(para.alignment) if para.alignment else "N/A"
    print(f"  段落{i:2d}: [{align}] {text_preview}")

print(f"\n[总段落数: {len(doc.paragraphs)}]")
print(f"\n{'=' * 60}")
print("验证完成")
print(f"{'=' * 60}")
