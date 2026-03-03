#!/usr/bin/env python3
"""
生成 2025年12月 警情分析研判报告
基于模板格式规范和提取的数据，智能仿写生成新报告
"""
import json
import copy
from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 1. 加载数据
# ============================================================
BASE_DIR = "/home/orangels/xm_dev/ls_dev/reportSkillMaker"
SESSION_DIR = os.path.join(BASE_DIR, "middle_file/1772532180011_session")
TEMPLATE_PATH = os.path.join(BASE_DIR, "关于11月份警情分析研判报告.docx")
OUTPUT_PATH = os.path.join(BASE_DIR, "output/output_2025年12月统计报告.docx")

with open(os.path.join(SESSION_DIR, "extracted_data.json"), "r", encoding="utf-8") as f:
    data = json.load(f)

# ============================================================
# 2. 使用模板作为基础，清空内容后重新写入
# ============================================================
# 我们采用"基于模板复制"的方式：
#   - 从原始模板复制文档，保留页面设置和section属性
#   - 清空所有段落内容
#   - 按格式规范重新写入新内容

doc = Document(TEMPLATE_PATH)

# ============================================================
# 3. 格式常量定义
# ============================================================
FONT_FANGSONG = "仿宋"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体"
FONT_FZXBSJT = "方正小标宋简体"

SIZE_55PT = Pt(55)    # 红头
SIZE_22PT = Pt(22)    # 主标题
SIZE_16PT = Pt(16)    # 正文/一级/二级标题
SIZE_14PT = Pt(14)    # 抄送

LINE_SPACING = 560  # 28pt 固定行距 = 560 twip (w:line 属性使用 twips 单位)
FIRST_LINE_INDENT = Emu(406400)   # 约2字符
FIRST_LINE_INDENT_2 = Emu(408305) # 三级标题用的缩进

COLOR_RED = RGBColor(0xFF, 0x00, 0x00)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)

# ============================================================
# 4. 辅助函数
# ============================================================

def set_paragraph_format(paragraph, alignment=None, first_line_indent=None,
                         left_indent=None, line_spacing=LINE_SPACING,
                         line_rule="exact", widow_control=False):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        # 直接通过 XML 设置行距，line_spacing 值为 twips 单位
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="{int(line_spacing)}" w:lineRule="exact"/>')
            pPr.append(spacing)
        else:
            spacing.set(qn('w:lineRule'), 'exact')
            spacing.set(qn('w:line'), str(int(line_spacing)))

    # Disable widow control
    if not widow_control:
        pPr = paragraph._element.get_or_add_pPr()
        wc = pPr.find(qn('w:widowControl'))
        if wc is None:
            wc = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>')
            pPr.append(wc)
        else:
            wc.set(qn('w:val'), '0')


def set_run_font(run, font_name, font_size, bold=None, color=None, kern=0):
    """设置 run 的字体格式"""
    run.font.name = font_name
    run.font.size = font_size
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    if bold is not None:
        run.font.bold = bold
        # Also set bCs for complex script bold
        bCs = rPr.find(qn('w:bCs'))
        if bold:
            if bCs is None:
                bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
                rPr.append(bCs)
        else:
            if bCs is not None:
                rPr.remove(bCs)

    if color is not None:
        run.font.color.rgb = color

    # Set kern
    kern_elem = rPr.find(qn('w:kern'))
    if kern_elem is None:
        kern_elem = parse_xml(f'<w:kern {nsdecls("w")} w:val="{kern}"/>')
        rPr.append(kern_elem)
    else:
        kern_elem.set(qn('w:val'), str(kern))


def add_paragraph(doc_obj, alignment=None, first_line_indent=None,
                  left_indent=None, line_spacing=LINE_SPACING):
    """添加一个新段落并设置格式"""
    p = doc_obj.add_paragraph()
    set_paragraph_format(p, alignment=alignment, first_line_indent=first_line_indent,
                         left_indent=left_indent, line_spacing=line_spacing)
    return p


def add_run_to_para(paragraph, text, font_name=FONT_FANGSONG, font_size=SIZE_16PT,
                    bold=None, color=None):
    """向段落添加一个 run"""
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size, bold=bold, color=color)
    return run


def clear_document(doc_obj):
    """清除文档中所有段落"""
    body = doc_obj.element.body
    # Remove all paragraphs
    for p in body.findall(qn('w:p')):
        body.remove(p)
    # Remove all tables
    for t in body.findall(qn('w:tbl')):
        body.remove(t)


def format_percent(value, abs_val=False):
    """格式化百分比，保留一位小数，去掉末尾.0"""
    if abs_val:
        value = abs(value)
    result = f"{value:.1f}"
    if result.endswith('.0'):
        result = result[:-2]
    return result


def trend_word(change_rate):
    """根据环比变化率返回上升/下降/持平"""
    if change_rate > 0:
        return "上升"
    elif change_rate < 0:
        return "下降"
    else:
        return "持平"


# ============================================================
# 5. 清除文档内容，保留格式设置
# ============================================================
clear_document(doc)

# ============================================================
# 6. 开始写入新内容
# ============================================================

# ------ 红头：发文单位名称 ------
p_header = doc.add_paragraph()
set_paragraph_format(p_header, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=None)
# 先添加一个空的仿宋run（与模板一致）
r0 = p_header.add_run("")
set_run_font(r0, FONT_FANGSONG, Pt(16))
# 红头文字 - 分成多个run以匹配模板的fitText效果
r1 = p_header.add_run("临高")
set_run_font(r1, FONT_FZXBSJT, SIZE_55PT, color=COLOR_RED, kern=0)
r2 = p_header.add_run("县公安局情报指挥中")
set_run_font(r2, FONT_FZXBSJT, SIZE_55PT, color=COLOR_RED, kern=0)
r3 = p_header.add_run("心")
set_run_font(r3, FONT_FZXBSJT, SIZE_55PT, color=COLOR_RED, kern=0)

# 在红头段落中插入红色分隔线（从模板复制绘图对象）
# 由于分隔线是复杂的VML/Drawing对象，我们通过XML直接插入
# 先读取模板中的分隔线XML
template_doc = Document(TEMPLATE_PATH)
template_p0 = template_doc.paragraphs[0]
# 从模板的第一个段落中找到 AlternateContent（包含分隔线）
mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
for element in template_p0._element:
    tag = element.tag.split('}')[1] if '}' in element.tag else element.tag
    if tag == 'r':
        # Check for drawing/pict in runs
        for child in element:
            child_tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if child_tag == 'drawing' or child_tag == 'pict':
                # Copy the entire run with drawing
                new_run_elem = copy.deepcopy(element)
                p_header._element.append(new_run_elem)
    # Check for AlternateContent at paragraph level
    if 'AlternateContent' in element.tag:
        ac_copy = copy.deepcopy(element)
        p_header._element.append(ac_copy)

# ------ 空行 ------
p_blank = doc.add_paragraph()
set_paragraph_format(p_blank, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=560)

# ------ 主标题 ------
p_title = doc.add_paragraph()
set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=560)
r_title1 = p_title.add_run("关于")
set_run_font(r_title1, FONT_FZXBSJT, SIZE_22PT, color=COLOR_BLACK, kern=0)
r_title2 = p_title.add_run("12月份警情分析的报告")
set_run_font(r_title2, FONT_FZXBSJT, SIZE_22PT, color=COLOR_BLACK, kern=0)

# ------ 标题后空行 ------
p_blank2 = doc.add_paragraph()
set_paragraph_format(p_blank2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)

# ============================================================
# 7. 一、整体情况
# ============================================================
overall = data["一、整体情况"]
categories = data["一、整体情况"]["各警情大类"]

# 一级标题
p_h1 = doc.add_paragraph()
set_paragraph_format(p_h1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_h1, "一、整体情况", FONT_HEITI, SIZE_16PT, color=COLOR_BLACK)

# 整体情况正文
p_overall = doc.add_paragraph()
set_paragraph_format(p_overall, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)

effective = overall["有效警情"]
harass = overall["骚扰警情"]

# 构建整体情况段
text_parts = []
# 有效警情总量
add_run_to_para(p_overall, f"12月1日至31日我局共接报", FONT_FANGSONG, SIZE_16PT)
add_run_to_para(p_overall, "有效警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_overall, f"{effective['本月']}起（", FONT_FANGSONG, SIZE_16PT)
add_run_to_para(p_overall, "不含骚扰警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_overall, f"{harass['本月']}起），环比{trend_word(effective['环比变化率'])}{format_percent(effective['环比变化率'], True)}%。", FONT_FANGSONG, SIZE_16PT)

# 各警情大类
cat_order = ["刑事警情", "治安警情", "交通警情", "纠纷警情", "群众紧急求助", "其他警情"]
add_run_to_para(p_overall, "其中", FONT_FANGSONG, SIZE_16PT)
for i, cat_name in enumerate(cat_order):
    cat_data = categories[cat_name]
    separator = "；" if i < len(cat_order) - 1 else "。"
    add_run_to_para(p_overall, f"{cat_name}", FONT_FANGSONG, SIZE_16PT, bold=True)
    add_run_to_para(p_overall, f"{cat_data['本月']}起，环比{trend_word(cat_data['环比变化率'])}{format_percent(cat_data['环比变化率'], True)}%{separator}",
                    FONT_FANGSONG, SIZE_16PT)

# ============================================================
# 8. 二、上升警情类别分布
# ============================================================
p_h2 = doc.add_paragraph()
set_paragraph_format(p_h2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_h2, "二、上升警情类别分布", FONT_HEITI, SIZE_16PT, color=COLOR_BLACK)

# 确定要重点分析的上升警情类型：交通警情、纠纷警情
# （群众紧急求助和其他警情虽然上升，但内容较杂，不作为重点分析类型）
# 根据数据和模板模式，选择: 交通警情(上升15.38%)、纠纷警情(上升8.08%)

section_number = 1  # 二级章节计数

# ------ （一）交通警情分析 ------
traffic = data["二、各警情大类详细数据"]["交通警情"]
traffic_accident = traffic["交通事故分析"]

p_sub = doc.add_paragraph()
set_paragraph_format(p_sub, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_sub, f"（一）交通警情分析", FONT_KAITI, SIZE_16PT, color=COLOR_BLACK)

# 概述段
p_desc = doc.add_paragraph()
set_paragraph_format(p_desc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_desc, "我局共接报", FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_desc, "交通警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_desc, f"{traffic['总量']['本月']}起，环比上升{format_percent(traffic['总量']['环比变化率'], True)}%。其中交通事故警情{traffic_accident['总量']['本月']}起，环比上升{format_percent(traffic_accident['总量']['环比变化率'], True)}%。",
                FONT_FANGSONG, SIZE_16PT)

# 1.从警情类别分析
p_dim1 = doc.add_paragraph()
set_paragraph_format(p_dim1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim1, "1.从警情类别分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

# 事故子类数据
accident_types = traffic_accident["按事故子类分布"]["本月"]
accident_types_last = traffic_accident["按事故子类分布"]["上月"]
# 建立上月数据映射
last_month_map = {item["name"]: item["count"] for item in accident_types_last}

# 取前几个主要类型
main_types = accident_types[:4]
text_parts = []
for i, t in enumerate(main_types):
    last_count = last_month_map.get(t["name"], 0)
    if last_count > 0:
        change = (t["count"] - last_count) / last_count * 100
        change_text = f"，环比{trend_word(change)}{format_percent(change, True)}%"
    else:
        change_text = ""
    if i == 0:
        text_parts.append(f"主要集中在{t['name']}{t['count']}起{change_text}")
    else:
        text_parts.append(f"{t['name']}{t['count']}起{change_text}")
add_run_to_para(p_dim1, "；".join(text_parts) + "。", FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 2.从发案时间分析
p_dim2 = doc.add_paragraph()
set_paragraph_format(p_dim2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim2, "2.从发案时间分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

# 时段分析
time_slots = traffic["时段分布"]["本月"]
# 找高发时段
peak_slots = time_slots[:4]  # 前四个时段
peak_text = f"交通事故时段分布特征显著，{peak_slots[0]['name']}最为集中，发生{peak_slots[0]['count']}起；其次{peak_slots[1]['name']}发生{peak_slots[1]['count']}起，{peak_slots[2]['name']}发生{peak_slots[2]['count']}起。上述时段均对应午间出行、下午及傍晚通勤高峰。"

# 工作日vs周末
wd = traffic["工作日vs周末"]
weekend_pct = round(wd["weekend"] / (wd["weekday"] + wd["weekend"]) * 100, 1)

add_run_to_para(p_dim2, "一是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_dim2, peak_text, FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_dim2, "二是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_dim2, f"周六日节假日我局共接报交通警情{wd['weekend']}起，占全部交通警情的{format_percent(weekend_pct)}%，工作日共接报{wd['weekday']}起，节假日交通出行需求增大是事故高发的重要因素。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 3.从交通违法分析
traffic_violation = traffic["交通违法分析"]
p_dim3 = doc.add_paragraph()
set_paragraph_format(p_dim3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim3, "3.从交通违法分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

violation_types = traffic_violation["按细类分布"]
viol_text_parts = []
for i, v in enumerate(violation_types[:4]):
    if i == 0:
        viol_text_parts.append(f"交通违法警情{traffic_violation['总量']['本月']}起，环比上升{format_percent(traffic_violation['总量']['环比变化率'], True)}%。主要集中在{v['name']}{v['count']}起（占比{format_percent(v['percentage'])}%）")
    else:
        viol_text_parts.append(f"{v['name']}{v['count']}起（占比{format_percent(v['percentage'])}%）")
add_run_to_para(p_dim3, "，其次".join(viol_text_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 小结
p_summary = doc.add_paragraph()
set_paragraph_format(p_summary, first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_summary, "小结：", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_summary, "近期辖区交通警情及交通事故警情环比均呈上升态势，上升幅度较为明显。此类警情特征突出：",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_summary, "一是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_summary, "机动车与机动车事故仍为主要事故类型，且环比上升显著，需持续加强路面巡查和交通秩序维护；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_summary, "二是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_summary, "事故高发时段集中在午间至傍晚通勤高峰期，与人流车流密集密切相关；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_summary, "三是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_summary, f"交通违法警情环比大幅上升{format_percent(traffic_violation['总量']['环比变化率'], True)}%，其中交通事故逃逸占比最高，需加大查处力度。建议交警部门针对性强化重点时段、重点路段管控，切实遏制交通事故上升势头。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# ------ （二）纠纷警情分析 ------
dispute = data["二、各警情大类详细数据"]["纠纷警情"]

p_sub2 = doc.add_paragraph()
set_paragraph_format(p_sub2, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_sub2, "（二）纠纷警情分析", FONT_KAITI, SIZE_16PT, color=COLOR_BLACK)

# 概述
p_desc2 = doc.add_paragraph()
set_paragraph_format(p_desc2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_desc2, "我局共接报", FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_desc2, "纠纷警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_desc2, f"{dispute['总量']['本月']}起，环比上升{format_percent(dispute['总量']['环比变化率'], True)}%。",
                FONT_FANGSONG, SIZE_16PT)

# 1.从纠纷类型分析
p_dim_d1 = doc.add_paragraph()
set_paragraph_format(p_dim_d1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim_d1, "1.从纠纷类型分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

disp_types = dispute["按报警类型分布"]["本月"]
top_disp = disp_types[:5]
disp_text_parts = []
for i, d in enumerate(top_disp):
    if i == 0:
        disp_text_parts.append(f"主要集中在{d['name']}{d['count']}起（占比{format_percent(d['percentage'])}%）")
    else:
        disp_text_parts.append(f"{d['name']}{d['count']}起（占比{format_percent(d['percentage'])}%）")
add_run_to_para(p_dim_d1, "，其次".join(disp_text_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 2.从发案时间分析
p_dim_d2 = doc.add_paragraph()
set_paragraph_format(p_dim_d2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim_d2, "2.从发案时间分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

disp_time = dispute["时段分布"]["本月"]
disp_wd = dispute["工作日vs周末"]
disp_weekend_pct = round(disp_wd["weekend"] / (disp_wd["weekday"] + disp_wd["weekend"]) * 100, 1)

add_run_to_para(p_dim_d2, f"纠纷警情时段分布以白天为主，{disp_time[0]['name']}最为集中共{disp_time[0]['count']}起（占比{format_percent(disp_time[0]['percentage'])}%），其次{disp_time[1]['name']}共{disp_time[1]['count']}起（占比{format_percent(disp_time[1]['percentage'])}%），{disp_time[2]['name']}共{disp_time[2]['count']}起（占比{format_percent(disp_time[2]['percentage'])}%）。工作日共接报{disp_wd['weekday']}起，周末共接报{disp_wd['weekend']}起，占比{format_percent(disp_weekend_pct)}%。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 3.从辖区分布分析
p_dim_d3 = doc.add_paragraph()
set_paragraph_format(p_dim_d3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_dim_d3, "3.从辖区分布分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

disp_area = dispute["按辖区分布"]["本月"]
area_parts = []
for i, a in enumerate(disp_area[:5]):
    if i == 0:
        area_parts.append(f"主要集中在{a['name']}{a['count']}起（占比{format_percent(a['percentage'])}%）")
    else:
        area_parts.append(f"{a['name']}{a['count']}起（占比{format_percent(a['percentage'])}%）")
add_run_to_para(p_dim_d3, "，其次".join(area_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 小结
p_sum_d = doc.add_paragraph()
set_paragraph_format(p_sum_d, first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sum_d, "小结：", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "近期纠纷警情呈小幅上升态势，矛盾纠纷风险不容忽视。此类警情特征突出：",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "一是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "纠纷类型多元，其他纠纷、产权权属纠纷、经济纠纷占比靠前，土地权属纠纷仍为细分类中最突出的矛盾点；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "二是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_d, f"噪音纠纷环比上升明显（本月{21}起，上月{13}起），年末社会活动增多是重要诱因；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "三是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_d, "辖区分布集中在西门所、东门所等城区派出所，需加强基层矛盾排查化解，推动纠纷多元调处机制，严防矛盾激化升级。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)


# ============================================================
# 9. 三、下降警情类型分布
# ============================================================
p_h3 = doc.add_paragraph()
set_paragraph_format(p_h3, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_h3, "三、下降警情类型分布", FONT_HEITI, SIZE_16PT, color=COLOR_BLACK)

# ------ （一）刑事警情分析 ------
criminal = data["二、各警情大类详细数据"]["刑事警情"]

p_sub3 = doc.add_paragraph()
set_paragraph_format(p_sub3, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_sub3, "（一）刑事警情分析", FONT_KAITI, SIZE_16PT, color=COLOR_BLACK)

# 概述
p_desc3 = doc.add_paragraph()
set_paragraph_format(p_desc3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_desc3, "我局共接报", FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_desc3, "刑事警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_desc3, f"{criminal['总量']['本月']}起，环比下降{format_percent(criminal['总量']['环比变化率'], True)}%。",
                FONT_FANGSONG, SIZE_16PT)

# 1.从警情类型分析
p_crim1 = doc.add_paragraph()
set_paragraph_format(p_crim1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_crim1, "1.从警情类型分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

crim_types = criminal["按报警类型分布"]["本月"]
crim_text_parts = []
for i, c in enumerate(crim_types):
    crim_text_parts.append(f"{c['name']}{c['count']}起（占比{format_percent(c['percentage'])}%）")
add_run_to_para(p_crim1, "，".join(crim_text_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 2.从细类分析
p_crim2 = doc.add_paragraph()
set_paragraph_format(p_crim2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_crim2, "2.从高发类型分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

crim_sub = criminal["按细类分布"]["本月"]
crim_sub_last = {item["name"]: item["count"] for item in criminal["按细类分布"]["上月"]}
crim_sub_parts = []
for i, c in enumerate(crim_sub[:4]):
    last_c = crim_sub_last.get(c["name"], 0)
    if last_c > 0:
        chg = (c["count"] - last_c) / last_c * 100
        chg_text = f"，较上月{last_c}起环比{trend_word(chg)}{format_percent(chg, True)}%"
    elif last_c == 0 and c["count"] > 0:
        chg_text = f"，上月{last_c}起"
    else:
        chg_text = ""
    if i == 0:
        crim_sub_parts.append(f"主要集中在{c['name']}{c['count']}起（占比{format_percent(c['percentage'])}%）{chg_text}")
    else:
        crim_sub_parts.append(f"{c['name']}{c['count']}起（占比{format_percent(c['percentage'])}%）{chg_text}")
add_run_to_para(p_crim2, "，其次".join(crim_sub_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 3.从辖区分布分析
p_crim3 = doc.add_paragraph()
set_paragraph_format(p_crim3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_crim3, "3.从辖区分布分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

crim_area = criminal["按辖区分布"]["本月"]
# Only list the top ones (not external jurisdictions)
local_areas = [a for a in crim_area if "市公安局" not in a["name"] and "情指中心" not in a["name"]]
area_crim_parts = []
for i, a in enumerate(local_areas[:5]):
    if i == 0:
        area_crim_parts.append(f"主要发生在{a['name']}{a['count']}起（占比{format_percent(a['percentage'])}%）")
    else:
        area_crim_parts.append(f"{a['name']}{a['count']}起（占比{format_percent(a['percentage'])}%）")
add_run_to_para(p_crim3, "，其次".join(area_crim_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 小结
p_sum_c = doc.add_paragraph()
set_paragraph_format(p_sum_c, first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sum_c, "小结：", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "近期刑事警情环比下降明显，打击成效初步显现。此类警情特征突出：",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "一是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "侵犯财产权利类仍为刑事警情主要类型，盗窃案件占比最高；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "二是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_c, f"电信网络诈骗{crim_sub[1]['count']}起，较上月{crim_sub_last.get('电信网络诈骗', 0)}起有所下降，但仍需保持高压态势；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "三是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_c, "案件辖区分布较分散，各派出所需继续深化严打整治，巩固降势成果，严防刑事警情反弹。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# ------ （二）治安警情分析 ------
security = data["二、各警情大类详细数据"]["治安警情"]

p_sub4 = doc.add_paragraph()
set_paragraph_format(p_sub4, first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_sub4, "（二）治安警情分析", FONT_KAITI, SIZE_16PT, color=COLOR_BLACK)

# 概述
p_desc4 = doc.add_paragraph()
set_paragraph_format(p_desc4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_desc4, "我局共接报", FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_desc4, "治安警情", FONT_FANGSONG, SIZE_16PT, bold=True)
add_run_to_para(p_desc4, f"{security['总量']['本月']}起，环比下降{format_percent(security['总量']['环比变化率'], True)}%。",
                FONT_FANGSONG, SIZE_16PT)

# 1.从警情类型分析
p_sec1 = doc.add_paragraph()
set_paragraph_format(p_sec1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sec1, "1.从警情类型分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

sec_types = security["按细类分布"]["本月"]
sec_types_last = {item["name"]: item["count"] for item in security["按细类分布"]["上月"]}
sec_parts = []
for i, s in enumerate(sec_types[:5]):
    last_s = sec_types_last.get(s["name"], 0)
    if last_s > 0:
        chg = (s["count"] - last_s) / last_s * 100
        chg_text = f"，环比{trend_word(chg)}{format_percent(chg, True)}%"
    else:
        chg_text = ""
    if i == 0:
        sec_parts.append(f"主要集中在{s['name']}{s['count']}起（占比{format_percent(s['percentage'])}%）{chg_text}")
    else:
        sec_parts.append(f"{s['name']}{s['count']}起（占比{format_percent(s['percentage'])}%）{chg_text}")
add_run_to_para(p_sec1, "，其次".join(sec_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 2.从涉刀警情分析
knife_data = data["三、专项数据"]["涉刀警情"]
p_sec2 = doc.add_paragraph()
set_paragraph_format(p_sec2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sec2, "2.从涉刀警情分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sec2, f"我局共接报涉刀警情{knife_data['总量']['本月']}起，较上月{knife_data['总量']['上月']}起环比上升{format_percent(knife_data['总量']['环比变化率'], True)}%。从警情大类分布看，",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
knife_by_type = knife_data["按警情大类分布"]
knife_parts = []
for k in knife_by_type[:4]:
    knife_parts.append(f"{k['name']}{k['count']}起（占比{format_percent(k['percentage'])}%）")
add_run_to_para(p_sec2, "、".join(knife_parts) + "。涉刀警情需高度关注，严防发生恶性案事件。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 3.从辖区分布分析
p_sec3 = doc.add_paragraph()
set_paragraph_format(p_sec3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sec3, "3.从辖区分布分析。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)

# 辖区分布 - 分殴打和盗窃两类
p_sec3_sub = doc.add_paragraph()
set_paragraph_format(p_sec3_sub, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sec3_sub, "（1）殴打警情。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
beat_area = security["殴打他人分析"]["按辖区分布"]
beat_parts = []
for i, b in enumerate(beat_area[:5]):
    if i == 0:
        beat_parts.append(f"主要发生在{b['name']}{b['count']}起（占比{format_percent(b['percentage'])}%）")
    else:
        beat_parts.append(f"{b['name']}{b['count']}起（占比{format_percent(b['percentage'])}%）")
add_run_to_para(p_sec3_sub, "，其次".join(beat_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

p_sec3_sub2 = doc.add_paragraph()
set_paragraph_format(p_sec3_sub2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sec3_sub2, "（2）盗窃警情。", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
theft_area = security["盗窃分析"]["按辖区分布"]
theft_parts = []
for i, t in enumerate(theft_area[:5]):
    if i == 0:
        theft_parts.append(f"主要发生在{t['name']}{t['count']}起（占比{format_percent(t['percentage'])}%）")
    else:
        theft_parts.append(f"{t['name']}{t['count']}起（占比{format_percent(t['percentage'])}%）")
add_run_to_para(p_sec3_sub2, "，其次".join(theft_parts) + "。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# 小结
p_sum_s = doc.add_paragraph()
set_paragraph_format(p_sum_s, first_line_indent=FIRST_LINE_INDENT_2, line_spacing=560)
add_run_to_para(p_sum_s, "小结：", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_s, "近期治安警情环比小幅下降，总体态势平稳可控。此类警情特征突出：",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_s, "一是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_s, f"盗窃和殴打他人仍为治安警情两大主要类型，分别占比{format_percent(sec_types[0]['percentage'])}%和{format_percent(sec_types[1]['percentage'])}%，需持续保持严打态势；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_s, "二是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_s, f"涉刀警情环比上升{format_percent(knife_data['总量']['环比变化率'], True)}%，风险不容忽视，需加强管制刀具收缴和涉刀警情快侦快办；",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)
add_run_to_para(p_sum_s, "三是", FONT_FANGSONG, SIZE_16PT, bold=True, color=COLOR_BLACK)
add_run_to_para(p_sum_s, "西门所、东门所仍为治安警情高发区域，需加强巡逻防控和见警率，有效遏制治安案件多发势头。",
                FONT_FANGSONG, SIZE_16PT, color=COLOR_BLACK)

# ============================================================
# 10. 四、工作建议
# ============================================================
p_h4 = doc.add_paragraph()
set_paragraph_format(p_h4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_h4, "四、工作建议", FONT_HEITI, SIZE_16PT, color=COLOR_BLACK)

# （一）强化交通安全管控
p_adv1 = doc.add_paragraph()
set_paragraph_format(p_adv1, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_adv1, "（一）强化路段时段管控。", FONT_KAITI, SIZE_16PT, bold=False, color=COLOR_BLACK)
add_run_to_para(p_adv1, f"交警部门要在每日14时至20时重点时段和节假日，对主干道、学校周边、商圈路口增派警力，加大对交通事故逃逸的查处力度，每月联合开展不少于2次集中整治行动；同时关注机动车与机动车事故占比偏高问题，对事故多发路段开展隐患排查，强化交通安全宣传教育，切实压降交通事故发生率。",
                FONT_FANGSONG, SIZE_16PT, bold=False, color=COLOR_BLACK)

# （二）深化矛盾纠纷排查化解
p_adv2 = doc.add_paragraph()
set_paragraph_format(p_adv2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_adv2, "（二）深化矛盾纠纷排查化解。", FONT_KAITI, SIZE_16PT, bold=False, color=COLOR_BLACK)
add_run_to_para(p_adv2, f"各派出所要联合镇村（社区）力量每月对产权权属纠纷、经济纠纷、噪音纠纷等高发镇村开展不少于1次集中排查，建立重点人员台账，推动司法调解前置，严防矛盾升级；西门所、东门所、博厚所等纠纷警情高发单位要增加社区走访频次，每周不少于2次深入矛盾集中区域开展排查化解工作。",
                FONT_FANGSONG, SIZE_16PT, bold=False, color=COLOR_BLACK)

# （三）严打整治巩固降势
p_adv3 = doc.add_paragraph()
set_paragraph_format(p_adv3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_adv3, "（三）严打整治巩固降势。", FONT_KAITI, SIZE_16PT, bold=False, color=COLOR_BLACK)
add_run_to_para(p_adv3, f"各派出所要持续深化严打整治行动，巩固刑事警情、治安警情环比下降的良好态势。针对盗窃电动车等高发类型，加强技防物防建设，强化重点时段巡逻防控；针对电信网络诈骗，持续深化反诈宣传进社区、进校园，每月开展不少于2次集中宣传活动，提升群众防骗意识。西门所、东门所要严格落实亮灯巡逻和见警率管事率考核要求，有效遏制治安案件反弹势头。",
                FONT_FANGSONG, SIZE_16PT, bold=False, color=COLOR_BLACK)

# （四）加强涉刀警情防控
p_adv4 = doc.add_paragraph()
set_paragraph_format(p_adv4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=FIRST_LINE_INDENT, line_spacing=560)
add_run_to_para(p_adv4, "（四）加强涉刀警情防控。", FONT_KAITI, SIZE_16PT, bold=False, color=COLOR_BLACK)
add_run_to_para(p_adv4, f"各派出所要高度重视涉刀警情环比上升{format_percent(knife_data['总量']['环比变化率'], True)}%的突出问题，加大管制刀具收缴力度，对涉刀警情做到快侦快办、依法严处。西门所、东门所、多文所等涉刀警情多发辖区要加强夜间巡逻力度，每周组织不少于2次涉刀隐患排查，强化重点人员管控，严防涉刀恶性案事件发生。",
                FONT_FANGSONG, SIZE_16PT, bold=False, color=COLOR_BLACK)

# ============================================================
# 11. 署名
# ============================================================
# 空行
for _ in range(3):
    p_blank = doc.add_paragraph()
    set_paragraph_format(p_blank, line_spacing=560)

# 署名单位
p_sign = doc.add_paragraph()
set_paragraph_format(p_sign, first_line_indent=Emu(2641600), line_spacing=560)
add_run_to_para(p_sign, "临高县公安局情报指挥中心", FONT_FANGSONG, SIZE_16PT)

# 署名日期
p_date = doc.add_paragraph()
set_paragraph_format(p_date, first_line_indent=Emu(3048000), line_spacing=560)
add_run_to_para(p_date, "2026年1月6日", FONT_FANGSONG, SIZE_16PT)

# ============================================================
# 12. 空行 + 分隔线 + 抄送/抄报 + 印发行
# ============================================================
# 添加足够的空行以确保抄送在新页面底部
for _ in range(10):
    p_blank = doc.add_paragraph()
    set_paragraph_format(p_blank, line_spacing=560)

# 抄送
p_copy = doc.add_paragraph()
set_paragraph_format(p_copy, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=560)
# Set left_indent and hanging indent like template
pf = p_copy.paragraph_format
pf.left_indent = Emu(380365)
pf.first_line_indent = Emu(-177800)
add_run_to_para(p_copy, "", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_copy, "抄送", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_copy, "：", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_copy, "各所、队、室（中心）", FONT_FANGSONG, SIZE_14PT)

# 抄报
p_report = doc.add_paragraph()
set_paragraph_format(p_report, line_spacing=560)
pf_r = p_report.paragraph_format
pf_r.first_line_indent = Emu(177800)
add_run_to_para(p_report, "抄报", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_report, "：", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_report, "严树勋副县长，各局领导", FONT_FANGSONG, SIZE_14PT)

# 印发行
p_print = doc.add_paragraph()
set_paragraph_format(p_print, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=560)
pf_pr = p_print.paragraph_format
pf_pr.first_line_indent = Emu(177800)
add_run_to_para(p_print, "", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_print, "", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_print, "临高县公安局", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_print, "情报指挥中心", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_print, "             ", FONT_FANGSONG, SIZE_14PT)
add_run_to_para(p_print, "2026年1月6日印发", FONT_FANGSONG, SIZE_14PT)

# ============================================================
# 13. 保存文档
# ============================================================
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"报告已生成: {OUTPUT_PATH}")
print("生成完毕!")
