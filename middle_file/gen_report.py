#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def add_run_with_font(paragraph, text, font_name='仿宋', font_size=16, bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run

def set_paragraph_format(p, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.33)
    p.paragraph_format.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)

# 创建文档
doc = Document()

# 文头:单位名称(红色 55pt 居中)
p = doc.add_paragraph()
set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_run_with_font(p, '临高县公安局情报指挥中心', '方正小标宋简体', 55, color=RGBColor(255, 0, 0))

# 标题(22pt 居中)
p = doc.add_paragraph()
set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_run_with_font(p, '关于12月份警情分析的报告', '方正小标宋简体', 22)

# 空行
doc.add_paragraph()

# 一、整体情况
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '一、整体情况', '黑体', 16)

# 整体情况内容
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月1日至31日我局共接报')
add_run_with_font(p, '有效警情', bold=True)
add_run_with_font(p, '2890起（')
add_run_with_font(p, '不含骚扰警情', bold=True)
add_run_with_font(p, '124起），环比上升11.8%。其中')
add_run_with_font(p, '刑事警情', bold=True)
add_run_with_font(p, '18起，环比下降25.0%；')
add_run_with_font(p, '治安警情', bold=True)
add_run_with_font(p, '157起，环比下降4.3%；')
add_run_with_font(p, '交通警情', bold=True)
add_run_with_font(p, '923起，环比上升15.4%；')
add_run_with_font(p, '纠纷警情', bold=True)
add_run_with_font(p, '281起，环比上升8.1%；')
add_run_with_font(p, '群众紧急求助', bold=True)
add_run_with_font(p, '858起，环比上升8.2%；')
add_run_with_font(p, '其他警情', bold=True)
add_run_with_font(p, '653起，环比上升19.8%。')

# 空行
doc.add_paragraph()

# 二、上升警情类别分布
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '二、上升警情类别分布', '黑体', 16)

# (一)交通警情
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（一）交通警情', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份交通警情共接报923起，环比上升15.4%，呈现明显上升态势。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '1.从警情类别分析。')
add_run_with_font(p, '机动车与机动车事故322起，环比上升23.8%，占比34.9%；机动车与非机动车事故140起，环比下降4.1%，占比15.2%；单方事故130起，环比上升13.0%，占比14.1%；非机动车与非机动车事故48起，环比上升6.7%，占比5.2%；其他道路交通事故36起，环比上升80.0%，占比3.9%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '2.从发案时间分析。')
add_run_with_font(p, '一是', bold=True)
add_run_with_font(p, '时段分布上，12-18时发案389起，占比42.1%，为事故高发时段；18-24时发案256起，占比27.7%；6-12时发案227起，占比24.6%；0-6时发案51起，占比5.5%。')
add_run_with_font(p, '二是', bold=True)
add_run_with_font(p, '午后至傍晚时段交通流量大，事故风险较高，需加强重点时段路面管控。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '3.从辖区分布分析。')
add_run_with_font(p, '临高县公安局交通管理大队接报907起，占比98.3%，为主要责任单位。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '小结：', bold=True)
add_run_with_font(p, '交通警情呈现明显上升态势，且上升幅度较大。此类警情特征突出：')
add_run_with_font(p, '一是', bold=True)
add_run_with_font(p, '机动车与机动车事故增幅最大，环比上升23.8%，反映出机动车保有量增加带来的交通压力；')
add_run_with_font(p, '二是', bold=True)
add_run_with_font(p, '午后至傍晚时段为事故高发期，12-18时占比超过四成，需加强重点时段路面管控和交通疏导；')
add_run_with_font(p, '三是', bold=True)
add_run_with_font(p, '交通管理大队承担了98.3%的警情处置任务，工作压力较大，需统筹警力资源，提升快速反应和处置能力。')

# (二)其他警情
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（二）其他警情', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份其他警情共接报653起，环比上升19.8%，上升幅度较大。此类警情涵盖范围广泛，包括群众咨询、求助等非紧急类警情，反映出群众对公安机关服务需求的增加。需进一步优化接处警流程，提升服务质量和效率。')

# (三)纠纷警情
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（三）纠纷警情', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份纠纷警情共接报281起，环比上升8.1%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '1.从纠纷类型分析。')
add_run_with_font(p, '土地权属纠纷39起，占比13.9%；拖欠工资17起，占比6.0%；殴打他人、故意伤害他人身体17起，占比6.0%；家庭纠纷15起，占比5.3%；其他经济纠纷12起，占比4.3%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '2.从辖区分布分析。')
add_run_with_font(p, '临高临城西门派出所接报76起，占比27.0%；临高临城东门派出所接报47起，占比16.7%；临高博厚海岸派出所接报25起，占比8.9%；临高马袅海岸派出所接报15起，占比5.3%；临高皇桐派出所接报13起，占比4.6%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '小结：', bold=True)
add_run_with_font(p, '纠纷警情呈现上升态势。此类警情特征突出：')
add_run_with_font(p, '一是', bold=True)
add_run_with_font(p, '土地权属纠纷占比最高，达13.9%，反映出农村地区土地矛盾依然突出，需加强源头治理和司法调解；')
add_run_with_font(p, '二是', bold=True)
add_run_with_font(p, '拖欠工资和经济纠纷环比上升明显，分别上升30.8%和33.3%，需关注劳资矛盾和经济纠纷的排查化解；')
add_run_with_font(p, '三是', bold=True)
add_run_with_font(p, '西门所和东门所警情占比较高，合计达43.7%，需加强重点区域的矛盾纠纷排查和调处工作。')

# (四)群众紧急求助
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（四）群众紧急求助', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份群众紧急求助共接报858起，环比上升8.2%。此类警情涵盖群众遇到的各类紧急情况，包括走失人员查找、紧急救助等，反映出群众对公安机关的信任和依赖。需进一步提升应急处置能力，确保群众求助得到及时有效回应。')

# 空行
doc.add_paragraph()

# 三、下降警情类别分布
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '三、下降警情类别分布', '黑体', 16)

# (一)刑事警情
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（一）刑事警情', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份刑事警情共接报18起，环比下降25.0%，降幅明显。这一成效得益于前期严打整治行动的持续开展和重点区域的精准防控，刑事案件得到有效遏制。需继续保持高压态势，巩固防控成果。')

# (二)治安警情
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（二）治安警情', bold=True)

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '12月份治安警情共接报157起，环比下降4.3%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '1.从警情类型分析。')
add_run_with_font(p, '盗窃49起，环比下降21.0%，占比31.2%；殴打他人、故意伤害他人身体47起，环比下降11.3%，占比29.9%；故意损毁财物19起，环比上升46.2%，占比12.1%；威胁人身安全9起，环比上升200.0%，占比5.7%；家庭暴力6起，环比上升20.0%，占比3.8%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '2.从辖区分布分析。')
add_run_with_font(p, '临高临城西门派出所接报50起，占比31.8%；临高临城东门派出所接报26起，占比16.6%；临高新盈海岸派出所接报12起，占比7.6%；临高马袅海岸派出所接报12起，占比7.6%；临高加来派出所接报9起，占比5.7%。')

p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '小结：', bold=True)
add_run_with_font(p, '治安警情总体呈现下降态势，管控成效明显。但需关注：')
add_run_with_font(p, '一是', bold=True)
add_run_with_font(p, '盗窃警情环比下降21.0%，降幅明显，反映出技防物防措施和打击力度取得实效；')
add_run_with_font(p, '二是', bold=True)
add_run_with_font(p, '故意损毁财物和威胁人身安全警情环比上升较大，分别上升46.2%和200.0%，虽然基数较小，但需关注矛盾激化风险；')
add_run_with_font(p, '三是', bold=True)
add_run_with_font(p, '西门所和东门所警情占比较高，合计达48.4%，需继续加强重点区域的治安防控和巡逻管控。')

# 空行
doc.add_paragraph()

# 四、工作建议
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '四、工作建议', '黑体', 16)

# (一)
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（一）强化交通安全管控。', bold=True)
add_run_with_font(p, '交通管理大队要针对12-18时事故高发时段，每日开展不少于1次的重点路段巡查管控，加强交通疏导和违法行为查处；要联合镇村力量，每月对事故多发路段开展不少于2次的安全隐患排查，及时消除安全隐患；要加强交通安全宣传教育，每月进入社区、学校、企业开展不少于3次的宣传活动，提升群众交通安全意识，有效遏制交通事故上升势头。')

# (二)
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（二）深化矛盾纠纷化解。', bold=True)
add_run_with_font(p, '各派出所要联合镇村（社区）力量每月对土地纠纷、劳资纠纷等高发区域开展不少于1次集中排查，建立重点矛盾台账，推动司法调解前置，严防矛盾升级；西门所、东门所、博厚所等纠纷警情高发单位要加强与司法所、人民调解委员会的协作配合，每周召开不少于1次的联席会议，及时化解矛盾纠纷，筑牢社会稳定防线。')

# (三)
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（三）巩固治安防控成果。', bold=True)
add_run_with_font(p, '各派出所要继续保持对盗窃、殴打他人等治安违法行为的高压严打态势，对重点区域、重点时段开展"亮灯巡逻"，每周不少于3次，提升见警率、管事率；要密切关注故意损毁财物、威胁人身安全等警情上升态势，加强重点人员管控和矛盾纠纷排查，严防发生极端案事件；要持续推进技防物防建设，每月督促指导辖区重点场所完善视频监控、防盗设施，不断提升治安防控水平。')

# (四)
p = doc.add_paragraph()
set_paragraph_format(p)
add_run_with_font(p, '（四）提升应急服务能力。', bold=True)
add_run_with_font(p, '指挥中心要进一步优化接处警流程，对群众紧急求助类警情实行快速响应机制，确保第一时间调度警力处置；各派出所要加强应急处突演练，每月开展不少于1次的实战化训练，提升快速反应和应急处置能力；要强化与消防、医疗、民政等部门的协作联动，建立健全应急联动机制，确保群众求助得到及时有效回应，不断提升群众安全感和满意度。')

# 空行
doc.add_paragraph()
doc.add_paragraph()

# 文尾
p = doc.add_paragraph()
set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_run_with_font(p, '临高县公安局情报指挥中心')

p = doc.add_paragraph()
set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_run_with_font(p, '2025年12月31日')

# 保存文档
doc.save('/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_2025年12月统计报告.docx')
print('报告生成成功！')
