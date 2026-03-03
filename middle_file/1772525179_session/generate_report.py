#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
警情分析报告生成脚本
根据模板分析和提取数据智能生成报告
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

def add_red_line(paragraph):
    """添加红色装饰线"""
    run = paragraph.add_run()
    # 简化处理：使用下划线模拟红色装饰线
    paragraph.paragraph_format.border_bottom = True

def create_paragraph_with_format(doc, text, font_name, font_size, alignment, bold=False, color=None, first_line_indent=0):
    """创建带格式的段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    # 设置字体
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    # 设置中文字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 设置颜色
    if color:
        run.font.color.rgb = color

    # 设置段落格式
    para.alignment = alignment
    para.paragraph_format.line_spacing = Pt(28)  # 560twips = 28pt

    # 设置首行缩进
    if first_line_indent > 0:
        para.paragraph_format.first_line_indent = Inches(first_line_indent / 72)  # 转换为英寸

    return para

def add_mixed_paragraph(doc, parts, font_name, font_size, alignment, first_line_indent=0):
    """添加混合格式段落（部分加粗）"""
    para = doc.add_paragraph()

    for text, bold in parts:
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold

        # 设置中文字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 设置段落格式
    para.alignment = alignment
    para.paragraph_format.line_spacing = Pt(28)

    if first_line_indent > 0:
        para.paragraph_format.first_line_indent = Inches(first_line_indent / 72)

    return para

def generate_report(data_file, output_file):
    """生成报告"""

    # 读取数据
    with open(data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 创建文档
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # 1. 发文单位（顶部）
    create_paragraph_with_format(
        doc,
        "临高县公安局情报指挥中心",
        "方正小标宋简体",
        55,
        WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(255, 0, 0)
    )

    # 2. 主标题
    create_paragraph_with_format(
        doc,
        f"关于{data['报告基本信息']['目标月份']}警情分析的报告",
        "方正小标宋简体",
        22,
        WD_ALIGN_PARAGRAPH.CENTER
    )

    # 3. 一、整体情况
    create_paragraph_with_format(
        doc,
        "一、整体情况",
        "黑体",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # 整体情况正文
    overall = data['一级数据_整体情况']
    parts = [
        (f"{data['报告基本信息']['统计时间范围']['起始日期']}至{data['报告基本信息']['统计时间范围']['结束日期']}，我局共接报", False),
        ("有效警情", True),
        (f"{overall['有效警情总量']['本期']}起（不含", False),
        ("骚扰警情", True),
        (f"{overall['骚扰警情数量']['本期']}起），环比上升{overall['有效警情总量']['环比']}%。其中，", False),
        ("刑事警情", True),
        (f"{overall['刑事警情']['本期']}起，环比下降{abs(overall['刑事警情']['环比'])}%；", False),
        ("治安警情", True),
        (f"{overall['治安警情']['本期']}起，环比下降{abs(overall['治安警情']['环比'])}%；", False),
        ("交通警情", True),
        (f"{overall['交通警情']['本期']}起，环比上升{overall['交通警情']['环比']}%；", False),
        ("纠纷警情", True),
        (f"{overall['纠纷警情']['本期']}起，环比上升{overall['纠纷警情']['环比']}%；", False),
        ("群众紧急求助", True),
        (f"{overall['群众紧急求助']['本期']}起，环比上升{overall['群众紧急求助']['环比']}%；", False),
        ("其他警情", True),
        (f"{overall['其他警情']['本期']}起，环比上升{overall['其他警情']['环比']}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 4. 二、上升警情类别分布
    create_paragraph_with_format(
        doc,
        "二、上升警情类别分布",
        "黑体",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （一）治安殴打他人警情分析
    create_paragraph_with_format(
        doc,
        "（一）治安殴打他人警情分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    # 治安殴打总量
    ouда = data['二级数据_治安殴打']
    parts = [
        (f"{data['报告基本信息']['目标月份']}，我局共接报", False),
        ("治安殴打他人警情", True),
        (f"{ouда['总量']['本期']}起，环比下降{abs(ouда['总量']['环比'])}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 5.从辖区分布分析
    create_paragraph_with_format(
        doc,
        "1.从辖区分布分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True
    )

    # 辖区分布数据
    if ouда['按辖区分布']:
        top3 = ouда['按辖区分布'][:3]
        parts = [
            ("主要集中在", False),
            (f"{top3[0]['派出所']}", True),
            (f"{top3[0]['数量']}起（占比{top3[0]['占比']}%），其次", False),
            (f"{top3[1]['派出所']}", True),
            (f"{top3[1]['数量']}起（占比{top3[1]['占比']}%）、", False),
            (f"{top3[2]['派出所']}", True),
            (f"{top3[2]['数量']}起（占比{top3[2]['占比']}%）。", False)
        ]
        add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 小结
    parts = [
        ("小结：", True),
        (f"本月治安殴打警情呈现下降态势，环比下降{abs(ouда['总量']['环比'])}%，", False),
        (f"主要集中在{top3[0]['派出所']}辖区。", False),
        ("此类警情特征突出：", False),
        ("一是", True),
        ("辖区分布相对集中，城区派出所警情占比较高；", False),
        ("二是", True),
        ("多因口角纠纷、邻里矛盾等引发；", False),
        ("三是", True),
        ("部分案件涉及刀具等管制器具，存在安全隐患。建议各派出所加强矛盾纠纷排查化解，强化重点人员管控，持续开展「晚安行动」，及时消除治安隐患。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # （二）涉未成人警情分析
    create_paragraph_with_format(
        doc,
        "（二）涉未成人警情分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    # 涉未成人总量
    minor = data['二级数据_涉未成人']
    parts = [
        (f"{data['报告基本信息']['目标月份']}，我局共接报", False),
        ("涉未成人警情", True),
        (f"{minor['总量']['本期']}起，环比下降{abs(minor['总量']['环比'])}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 1.从警情类型分析
    create_paragraph_with_format(
        doc,
        "1.从警情类型分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True
    )

    type_data = minor['按警情类型分类']
    parts = [
        ("主要为", False),
        ("求助警情", True),
        (f"{type_data['求助警情']}起（占比{type_data['求助警情占比']}%），其次为", False),
        ("其他警情", True),
        (f"{type_data['其他警情']}起、", False),
        ("治安警情", True),
        (f"{type_data['治安警情']}起、", False),
        ("纠纷警情", True),
        (f"{type_data['纠纷警情']}起。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 2.从辖区分布分析
    create_paragraph_with_format(
        doc,
        "2.从辖区分布分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True
    )

    if minor['按辖区分布']:
        top2 = minor['按辖区分布'][:2]
        parts = [
            ("主要集中在", False),
            (f"{top2[0]['派出所']}", True),
            (f"{top2[0]['数量']}起（占比{top2[0]['占比']}%），其次", False),
            (f"{top2[1]['派出所']}", True),
            (f"{top2[1]['数量']}起（占比{top2[1]['占比']}%）。", False)
        ]
        add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 小结
    parts = [
        ("小结：", True),
        (f"本月涉未成人警情呈现下降态势，环比下降{abs(minor['总量']['环比'])}%，", False),
        ("以求助警情为主，占比超过三成。", False),
        ("此类警情特征突出：", False),
        ("一是", True),
        ("求助类警情占比较高，反映未成年人自我保护意识增强；", False),
        ("二是", True),
        (f"辖区分布集中在{top2[0]['派出所']}等城区派出所；", False),
        ("三是", True),
        ("涉及同学间打架、未成年人怀孕等敏感问题，需重点关注。建议各派出所持续开展「护苗行动」，加强校园周边巡逻防控，深化警校联动机制，及时发现和化解涉未成年人矛盾纠纷。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # （三）金牌港重点园区警情分析
    create_paragraph_with_format(
        doc,
        "（三）金牌港重点园区警情分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    # 金牌港总量
    park = data['二级数据_金牌港园区']
    parts = [
        (f"{data['报告基本信息']['目标月份']}，", False),
        ("金牌港重点园区", True),
        (f"共接报警情{park['总量']['本期']}起，环比上升{park['总量']['环比']}%。其中，", False),
        ("治安警情", True),
        (f"{park['治安警情']['本期']}起，环比下降{abs(park['治安警情']['环比'])}%；", False),
        ("纠纷警情", True),
        (f"{park['纠纷警情']['本期']}起；", False),
        ("紧急求助警情", True),
        (f"{park['紧急求助警情']['本期']}起，环比上升{park['紧急求助警情']['环比']}%；", False),
        ("其他警情", True),
        (f"{park['其他警情']['本期']}起，环比下降{abs(park['其他警情']['环比'])}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 小结
    parts = [
        ("小结：", True),
        (f"本月金牌港园区警情呈现上升态势，环比上升{park['总量']['环比']}%，", False),
        ("主要为紧急求助警情和纠纷警情。", False),
        ("此类警情特征突出：", False),
        ("一是", True),
        ("园区警情总量上升，反映园区治安形势需持续关注；", False),
        ("二是", True),
        ("纠纷警情从无到有，劳资纠纷、噪音纠纷等问题凸显；", False),
        ("三是", True),
        ("紧急求助警情大幅上升，需加强园区应急处置能力。建议加强园区日常巡逻防控，深化警企联动机制，及时排查化解矛盾纠纷，守护园区安全稳定。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 三、下降警情类型分布
    create_paragraph_with_format(
        doc,
        "三、下降警情类型分布",
        "黑体",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （一）交通事故警情分析
    create_paragraph_with_format(
        doc,
        "（一）交通事故警情分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    # 交通事故总量
    traffic = data['二级数据_交通事故']
    parts = [
        (f"{data['报告基本信息']['目标月份']}，我局共接报", False),
        ("交通警情", True),
        (f"{traffic['交通警情总量']['本期']}起，环比上升{traffic['交通警情总量']['环比']}%。其中，", False),
        ("交通事故警情", True),
        (f"{traffic['交通事故警情总量']['本期']}起，环比上升{traffic['交通事故警情总量']['环比']}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 1.从警情类别分析
    create_paragraph_with_format(
        doc,
        "1.从警情类别分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True
    )

    parts = [
        ("主要为", False),
        ("机动车与机动车事故", True),
        (f"{traffic['机动车与机动车事故']['本期']}起（环比上升{traffic['机动车与机动车事故']['环比']}%），其次为", False),
        ("机动车与非机动车事故", True),
        (f"{traffic['机动车与非机动车事故']['本期']}起（环比上升{traffic['机动车与非机动车事故']['环比']}%）、", False),
        ("单方事故", True),
        (f"{traffic['单方事故']['本期']}起（环比上升{traffic['单方事故']['环比']}%）。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 2.从发案时间分析
    create_paragraph_with_format(
        doc,
        "2.从发案时间分析",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True
    )

    time_data = traffic['按发案时间分析']
    parts = [
        ("主要集中在", False),
        ("16时至19时", True),
        (f"{time_data['16时至19时']}起，其次为", False),
        ("11时至14时", True),
        (f"{time_data['11时至14时']}起。", False),
        ("周六日节假日", True),
        (f"发生交通事故{time_data['周六日节假日']}起，占比{time_data['周六日节假日占比']}%。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 小结
    parts = [
        ("小结：", True),
        (f"本月交通事故警情呈现上升态势，环比上升{traffic['交通事故警情总量']['环比']}%，", False),
        ("机动车与机动车事故占比最高。", False),
        ("此类警情特征突出：", False),
        ("一是", True),
        ("16时至19时、11时至14时为事故高发时段，与交通出行高峰期吻合；", False),
        ("二是", True),
        ("周末节假日事故占比超过四分之一，反映节假日出行增多；", False),
        ("三是", True),
        ("机动车与机动车事故大幅上升，需加强路面管控。建议交警部门强化重点路段、重点时段巡逻管控，加大交通违法行为查处力度，深化交通安全宣传教育，有效预防和减少交通事故发生。", False)
    ]
    add_mixed_paragraph(doc, parts, "仿宋", 16, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

    # 四、工作建议
    create_paragraph_with_format(
        doc,
        "四、工作建议",
        "黑体",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （一）严打整治强化震慑
    create_paragraph_with_format(
        doc,
        "（一）严打整治强化震慑",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    create_paragraph_with_format(
        doc,
        "各派出所要持续开展「晚安行动」，加强夜间巡逻防控，每周不少于2次集中清查行动，重点打击殴打他人、寻衅滋事等违法犯罪行为，形成强大震慑。对涉刀警情要从严从快处置，加强管制器具收缴，及时消除安全隐患，有效压降治安殴打警情。",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （二）多举措防控筑牢安全屏障
    create_paragraph_with_format(
        doc,
        "（二）多举措防控筑牢安全屏障",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    create_paragraph_with_format(
        doc,
        "各派出所要深入开展「护苗行动」，加强校园周边巡逻防控，每月不少于1次进校园开展安全教育，深化警校联动机制。对涉未成年人矛盾纠纷要及时发现、及时化解，对涉未成年人违法犯罪要依法从严打击，切实保护未成年人合法权益，守护未成年人健康成长。",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （三）严管严控护航园区
    create_paragraph_with_format(
        doc,
        "（三）严管严控护航园区",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    create_paragraph_with_format(
        doc,
        "相关派出所要加强金牌港重点园区日常巡逻防控，深化警企联动机制，每月不少于1次进园区开展矛盾纠纷排查，及时发现和化解劳资纠纷、噪音纠纷等矛盾问题。要强化应急处置能力建设，提升快速反应水平，确保园区安全稳定，为园区发展营造良好治安环境。",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # （四）强化路段时段管控
    create_paragraph_with_format(
        doc,
        "（四）强化路段时段管控",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        first_line_indent=32
    )

    create_paragraph_with_format(
        doc,
        "交警部门要强化重点路段、重点时段巡逻管控，特别是16时至19时、11时至14时等事故高发时段，加大路面巡逻密度，每日不少于3次定点执勤。要加大交通违法行为查处力度，深化交通安全宣传教育，特别是针对周末节假日出行高峰，提前发布安全提示，有效预防和减少交通事故发生。",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=32
    )

    # 落款
    doc.add_paragraph()  # 空行
    create_paragraph_with_format(
        doc,
        "临高县公安局情报指挥中心",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.CENTER
    )

    create_paragraph_with_format(
        doc,
        data['报告基本信息']['落款日期'],
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.CENTER
    )

    doc.add_paragraph()  # 空行

    # 抄送抄报
    create_paragraph_with_format(
        doc,
        "> 抄送：各所、队、室（中心）",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.LEFT
    )

    create_paragraph_with_format(
        doc,
        "抄报：严树勋副县长，各局领导",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.LEFT
    )

    create_paragraph_with_format(
        doc,
        f"临高县公安局情报指挥中心 {data['报告基本信息']['落款日期']}印发",
        "仿宋",
        16,
        WD_ALIGN_PARAGRAPH.LEFT
    )

    # 保存文档
    doc.save(output_file)
    print(f"报告已生成：{output_file}")

if __name__ == "__main__":
    data_file = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/middle_file/1772525179_session/extracted_data.json"
    output_file = "/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_12月统计报告.docx"

    generate_report(data_file, output_file)
