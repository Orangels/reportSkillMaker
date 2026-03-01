#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能仿写生成2025年12月警情分析报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json

# 读取数据
with open('/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/extracted_data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 创建文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = '仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.styles['Normal'].font.size = Pt(16)

# 1. 发文单位(红色大标题)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('临高县公安局情报指挥中心')
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
run.font.size = Pt(55)
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

# 2. 主标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('关于12月份警情分析的报告')
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
run.font.size = Pt(22)
run.bold = True

# 3. 一、整体情况
p = doc.add_paragraph()
run = p.add_run('一、整体情况')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)
run.bold = True

# 整体情况段落
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)

# 计算环比
total_change_rate = round((data['overall_stats']['valid_cases_dec'] - data['overall_stats']['valid_cases_nov']) / data['overall_stats']['valid_cases_nov'] * 100, 1)

# 整体情况内容
text = f"{data['report_info']['period']}我局共接报"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('有效警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['overall_stats']['valid_cases_dec']}起（"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('不含骚扰警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['overall_stats']['harassment_cases_dec']}起），环比上升{total_change_rate}%。其中"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 六大类警情
categories = data['six_categories']
for i, (cat_name, cat_data) in enumerate(categories.items()):
    run = p.add_run(cat_name)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(16)
    run.bold = True
    
    change_text = "上升" if cat_data['change'] > 0 else "下降"
    text = f"{cat_data['dec_count']}起，环比{change_text}{abs(cat_data['change_rate'])}%"
    if i < len(categories) - 1:
        text += "；"
    else:
        text += "。"
    
    run = p.add_run(text)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(16)


# 二、上升警情类别分布
p = doc.add_paragraph()
run = p.add_run('二、上升警情类别分布')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)
run.bold = True

# (一)交通警情分析
p = doc.add_paragraph()
run = p.add_run('(一)交通警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('交通警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['six_categories']['交通警情']['dec_count']}起，环比大幅上升{data['six_categories']['交通警情']['change_rate']}%。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 1.从警情类别分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('1.从警情类别分析。')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

traffic_subtypes = data['traffic_detail']['subtypes']
top_3_traffic = sorted(traffic_subtypes.items(), key=lambda x: x[1]['dec_count'], reverse=True)[:3]

text = f"主要集中在{top_3_traffic[0][0]}{top_3_traffic[0][1]['dec_count']}起，"
text += f"其次{top_3_traffic[1][0]}{top_3_traffic[1][1]['dec_count']}起，"
text += f"{top_3_traffic[2][0]}{top_3_traffic[2][1]['dec_count']}起。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 2.从发案时间分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('2.从发案时间分析。')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

peak_hours = data['time_distribution']['peak_hours']
text = f"交通警情时段分布特征明显，{peak_hours[0]['hour_range']}最为集中，累计发生{peak_hours[0]['count']}起；"
text += f"其次{peak_hours[1]['hour_range']}，累计发生{peak_hours[1]['count']}起，"
text += f"两个时段均对应上午出行、下午通勤高峰。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月交通警情环比大幅上升，增幅显著。此类警情特征突出："
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('一是')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "机动车事故类型集中，机动车与机动车、机动车与非机动车事故占比较高，反映出路面车流量增大带来的安全风险；"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('二是')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "时段分布特征明显，集中在上午和下午通勤高峰时段，需针对性加强重点时段路面管控；"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('三是')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "年末岁尾人流车流密集，交通安全形势严峻，需持续强化交通安全宣传和路面执法力度。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)


# (二)涉刀警情分析
p = doc.add_paragraph()
run = p.add_run('(二)涉刀警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('涉刀警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['knife_cases']['total']}起，环比大幅上升。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 1.从警情类型分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('1.从警情类型分析。')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

knife_by_cat = data['knife_cases']['by_category']
top_knife_cats = sorted(knife_by_cat.items(), key=lambda x: x[1], reverse=True)[:3]
text = f"{top_knife_cats[0][0]}{top_knife_cats[0][1]}起，{top_knife_cats[1][0]}{top_knife_cats[1][1]}起，{top_knife_cats[2][0]}{top_knife_cats[2][1]}起。"
text += "其中治安警情占比最高，主要为殴打他人、故意伤害案件。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 2.从辖区分布分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('2.从辖区分布分析。')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

knife_by_jur = data['knife_cases']['by_jurisdiction']
top_knife_jur = sorted(knife_by_jur.items(), key=lambda x: x[1], reverse=True)[:3]
text = f"主要发生在{top_knife_jur[0][0]}{top_knife_jur[0][1]}起，{top_knife_jur[1][0]}{top_knife_jur[1][1]}起，{top_knife_jur[2][0]}{top_knife_jur[2][1]}起。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月涉刀警情环比大幅上升，安全风险突出。此类警情多发生在西门所、东门所等城区派出所辖区，"
text += "且多与治安殴打、故意伤害案件相关，造成当事人不同程度人身伤害。"
text += "需严格落实持刀人住所清查部署，加强重点人员管控，依法从严惩处涉刀违法犯罪，有效遏制此类警情上升势头。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)


# (三)涉未成人警情分析
p = doc.add_paragraph()
run = p.add_run('(三)涉未成人警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('涉未成人警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['minor_cases']['total']}起，环比显著上升。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 1.从警情类型分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('1.从警情类型分析。')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

minor_by_cat = data['minor_cases']['by_category']
top_minor_cats = sorted(minor_by_cat.items(), key=lambda x: x[1], reverse=True)[:3]
text = f"主要集中在{top_minor_cats[0][0]}{top_minor_cats[0][1]}起，"
text += f"其次{top_minor_cats[1][0]}{top_minor_cats[1][1]}起，"
text += f"{top_minor_cats[2][0]}{top_minor_cats[2][1]}起。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月涉未成人警情环比显著上升，其他警情和群众紧急求助占比较高。"
text += "年末岁尾未成年人安全风险增加，需持续强化护苗行动、法治宣讲校园行，"
text += "加强晚安守护巡逻，精准防范处置，全力守护未成年人健康成长。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# (四)金牌港重点园区警情分析
p = doc.add_paragraph()
run = p.add_run('(四)金牌港重点园区警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报涉金牌港重点园区警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

text = f"{data['jinpaigang_cases']['total']}起，环比上升。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 从警情类型分析
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('从警情类型分析，')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

jpg_by_cat = data['jinpaigang_cases']['by_category']
top_jpg_cats = sorted(jpg_by_cat.items(), key=lambda x: x[1], reverse=True)[:3]
text = f"主要集中在{top_jpg_cats[0][0]}{top_jpg_cats[0][1]}起，"
text += f"其次{top_jpg_cats[1][0]}{top_jpg_cats[1][1]}起，"
text += f"{top_jpg_cats[2][0]}{top_jpg_cats[2][1]}起。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月金牌港重点园区警情环比上升，交通警情占比最高，反映出园区车流量增大带来的管理压力。"
text += "需针对性强化园区交通管理和安全防范，优化接处警流程，全力维护园区治安秩序稳定。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)


# 三、下降警情类型分布
p = doc.add_paragraph()
run = p.add_run('三、下降警情类型分布')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)
run.bold = True

# (一)刑事警情分析
p = doc.add_paragraph()
run = p.add_run('(一)刑事警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('刑事警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['six_categories']['刑事警情']['dec_count']}起，环比下降{abs(data['six_categories']['刑事警情']['change_rate'])}%。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月刑事警情环比大幅下降，管控成效显著，需持续保持严打高压态势，巩固防控成果。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# (二)治安警情分析
p = doc.add_paragraph()
run = p.add_run('(二)治安警情分析')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('我局共接报')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

run = p.add_run('治安警情')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = f"{data['six_categories']['治安警情']['dec_count']}起，环比下降{abs(data['six_categories']['治安警情']['change_rate'])}%。"
run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 小结
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
run = p.add_run('小结：')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)
run.bold = True

text = "本月治安警情环比呈下降态势，治安防控工作成效明显，需持续强化重点区域巡逻防控，巩固良好态势。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)


# 四、工作建议
p = doc.add_paragraph()
run = p.add_run('四、工作建议')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)
run.bold = True

# (一)强化交通安全管控
p = doc.add_paragraph()
run = p.add_run('(一)强化交通安全管控')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
text = "交警部门要针对本月交通警情大幅上升态势，在每日10时至12时、14时至15时重点时段，"
text += "对主干道、学校周边、商圈路口增派警力，加强路面巡逻管控；"
text += "强化交通安全宣传引导，提升驾驶员安全意识；"
text += "严查交通违法行为，依法从严处罚，全力压降交通警情及交通事故。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# (二)严打涉刀违法犯罪
p = doc.add_paragraph()
run = p.add_run('(二)严打涉刀违法犯罪')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
text = "西门所、东门所等涉刀警情高发单位要严格落实持刀人住所清查部署工作要求，"
text += "建立重点人员台账，加强日常管控；"
text += "对持刀伤人案件快侦快办、依法从严惩处，"
text += "在辖区形成'带刀必查、涉刀必罚'的有力震慑，"
text += "切实压降涉刀警情，有效遏制此类警情上升势头。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# (三)多举措守护未成年人安全
p = doc.add_paragraph()
run = p.add_run('(三)多举措守护未成年人安全')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
text = "持续强化涉未成人法制宣传教育，全覆盖普及法律知识与自我防护技能；"
text += "各派出所要严格落实'晚安行动'要求，组织民辅警对网吧、河边、公园等未成年人易聚集区域开展巡查；"
text += "严厉打击各类涉未成人违法犯罪，形成有力震慑，切实守护未成年人身心健康与安全成长。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# (四)优化园区治安管理
p = doc.add_paragraph()
run = p.add_run('(四)优化园区治安管理')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
text = "马袅海岸派出所要紧盯金牌重点园区警情变化，聚焦交通警情占比较高问题，"
text += "深化警情研判预警，精准把握防控重点；"
text += "加强园区交通秩序管理，优化接处警流程、提升处置效能。"
text += "全力压降警情总量、防范风险隐患，切实维护金牌重点园区治安秩序稳定。"

run = p.add_run(text)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 落款
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('临高县公安局情报指挥中心')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(data['report_info']['report_date'])
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 抄送抄报
p = doc.add_paragraph()
run = p.add_run('抄送：各所、队、室(中心)')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

p = doc.add_paragraph()
run = p.add_run('抄报：严树勋副县长，各局领导')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

p = doc.add_paragraph()
run = p.add_run(f'临高县公安局情报指挥中心 {data["report_info"]["report_date"]}印发')
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(16)

# 保存文档
output_path = '/home/orangels/xm_dev/ls_dev/reportSkillMaker/output/output_2025年12月统计报告.docx'
doc.save(output_path)
print(f"报告已生成: {output_path}")

